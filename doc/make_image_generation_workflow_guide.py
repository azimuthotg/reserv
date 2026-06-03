from pathlib import Path
from textwrap import wrap

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parent
MD_PATH = ROOT / "image-generation-workflow-guide.md"
DOCX_PATH = ROOT / "image-generation-workflow-guide.docx"
PNG_PATH = ROOT / "image-generation-workflow-infographic.png"
FONT_REGULAR = Path(r"C:\Windows\Fonts\TH Sarabun New.ttf")
FONT_BOLD = Path(r"C:\Windows\Fonts\TH Sarabun New Bold.ttf")
FONT_ALT = Path(r"C:\Windows\Fonts\tahoma.ttf")
FONT_ALT_BOLD = Path(r"C:\Windows\Fonts\tahomabd.ttf")


COMMANDS = [
    ("START JOB", "เปิดงานใหม่ อ่าน context / progress / git"),
    ("CHECK ASSETS", "ตรวจรูป โลโก้ QR และ reference"),
    ("MAKE BRIEF", "สรุปเป้าหมาย ข้อความ mood/tone และข้อห้าม"),
    ("DRAFT LOW", "ทำร่างเร็วหรือ wireframe ดูโครงก่อน"),
    ("LAYOUT DRAFT", "วางข้อความจริง โลโก้จริง QR จริง"),
    ("REVIEW DRAFT", "ตรวจจุดดี จุดเสี่ยง และสิ่งที่ต้องแก้"),
    ("FINAL RENDER", "ทำภาพคุณภาพสูงหลัง layout ผ่านแล้ว"),
    ("EXPORT", "สร้างไฟล์ print/web และตรวจ QR"),
    ("RECORD JOB", "บันทึก progress, path และ commit"),
    ("END JOB", "สรุปผล ไฟล์ที่ได้ และสิ่งที่ยังค้าง"),
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_run_font(run, size=16, bold=False):
    run.font.name = "TH Sarabun New"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "TH Sarabun New")
    run.font.size = Pt(size)
    run.bold = bold


def add_paragraph(doc, text="", style=None, size=16, bold=False):
    paragraph = doc.add_paragraph(style=style)
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return paragraph


def add_bullets(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        run = paragraph.add_run(item)
        set_run_font(run, size=16)


def build_docx():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("คู่มือคำสั่งงานร่วมกัน\nสำหรับการผลิตป้ายและเจนภาพ")
    set_run_font(run, size=28, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Smart Creative Learning Space — สำนักวิทยบริการ มหาวิทยาลัยนครพนม")
    set_run_font(run, size=18)

    add_paragraph(doc, "หลักคิด", size=22, bold=True)
    add_paragraph(
        doc,
        "ไม่ข้ามจากไอเดียไป final render ทันที ให้เริ่มจาก brief, draft ต่ำ, layout draft, review แล้วค่อย render จริง",
        size=16,
    )

    add_paragraph(doc, "คำสั่งหลัก", size=22, bold=True)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "คำสั่ง"
    hdr[1].text = "ใช้เมื่อ"
    hdr[2].text = "ผลลัพธ์"
    for cell in hdr:
        set_cell_shading(cell, "1F4E79")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                set_run_font(run, size=15, bold=True)
                run.font.color.rgb = None

    rows = [
        ("START JOB", "เปิดงานใหม่", "อ่าน context, progress, git และสรุปขอบเขต"),
        ("CHECK ASSETS", "ตรวจวัตถุดิบ", "รู้ว่าใช้ไฟล์ไหนเป็นฉาก ตัวละคร โลโก้ QR"),
        ("MAKE BRIEF", "ก่อนออกแบบ", "ได้ brief พร้อม mood/tone และข้อห้าม"),
        ("DRAFT LOW", "ต้องการร่างเร็ว", "เห็นโครงสร้างโดยใช้ resource ต่ำ"),
        ("LAYOUT DRAFT", "ดูป้ายรวม", "มีข้อความจริง โลโก้จริง QR จริง"),
        ("REVIEW DRAFT", "ตรวจร่าง", "ได้จุดดี จุดเสี่ยง และรายการแก้"),
        ("FINAL RENDER", "หลังอนุมัติ layout", "ได้ภาพคุณภาพสูงตาม brief"),
        ("EXPORT", "ส่งใช้งาน", "ได้ไฟล์ print/web และตรวจ QR"),
        ("RECORD JOB", "บันทึกงาน", "อัปเดตเอกสาร progress และ path"),
        ("END JOB", "ปิดงาน", "สรุปไฟล์ commit และสิ่งค้าง"),
    ]
    for command, when, result in rows:
        cells = table.add_row().cells
        for idx, text in enumerate([command, when, result]):
            cells[idx].text = text
            for paragraph in cells[idx].paragraphs:
                for run in paragraph.runs:
                    set_run_font(run, size=14, bold=(idx == 0))

    add_paragraph(doc, "กติกาสำคัญ", size=22, bold=True)
    add_bullets(
        doc,
        [
            "ห้ามข้ามจาก brief ไป final render หากยังไม่ได้ตรวจ DRAFT LOW หรือ LAYOUT DRAFT",
            "QR Code และโลโก้ต้องใช้ไฟล์จริง ห้ามให้ AI สร้างเอง",
            "ข้อความสำคัญต้องวางด้วย layout tool ไม่ฝังในภาพ AI",
            "ภาพ AI ใช้เป็นภาพประกอบ ไม่ใช่แหล่งความจริงของข้อความหรือ QR",
            "ทุกไฟล์ที่ใช้งานจริงต้องเก็บ path ชัดเจนใน workspace",
        ],
    )

    add_paragraph(doc, "คัมภีร์การเจนภาพ", size=22, bold=True)
    add_paragraph(doc, "ให้ AI ทำ", size=18, bold=True)
    add_bullets(doc, ["ฉากห้องสมุด", "ตัวละครการ์ตูน", "mood/tone", "ภาพฐานสำหรับประกอบป้าย"])
    add_paragraph(doc, "ไม่ให้ AI ทำ", size=18, bold=True)
    add_bullets(doc, ["QR Code", "โลโก้", "ข้อความไทยยาว ๆ", "ตัดสินข้อมูลนโยบายเอง"])

    add_paragraph(doc, "Theme กลางของสำนักวิทยบริการ", size=22, bold=True)
    add_bullets(
        doc,
        [
            "ใช้ doc/illustrations/example1.jpg เป็น mood/tone reference เท่านั้น ไม่ลอก layout",
            "สีหลัก: น้ำเงินเข้ม, ทอง/เหลือง, ขาว/ครีม, ฟ้าอ่อน",
            "อารมณ์: สดใส เป็นมิตร น่ารัก สุภาพ",
            "personal_akeky.png ใช้เป็นมุม support ขนาดเล็ก ไม่แย่งความเด่นจาก QR และข้อความหลัก",
        ],
    )

    doc.add_page_break()
    add_paragraph(doc, "Infographic สรุป workflow", size=22, bold=True)
    doc.add_picture(str(PNG_PATH), width=Inches(6.8))
    doc.save(DOCX_PATH)


def font(path, size, fallback, bold=False):
    candidate = path if path.exists() else fallback
    return ImageFont.truetype(str(candidate), size=size)


def rounded(draw, xy, radius, fill, outline=None, width=2):
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=width)


def draw_wrapped(draw, text, xy, font_obj, fill, max_width, line_gap=6):
    x, y = xy
    line = ""
    for token in text.split(" "):
        test = token if not line else f"{line} {token}"
        if draw.textbbox((0, 0), test, font=font_obj)[2] <= max_width:
            line = test
        else:
            draw.text((x, y), line, font=font_obj, fill=fill)
            y += font_obj.size + line_gap
            line = token
    if line:
        draw.text((x, y), line, font=font_obj, fill=fill)
        y += font_obj.size + line_gap
    return y


def build_infographic():
    W, H = 1600, 2200
    navy = "#0b3f82"
    navy2 = "#082b5f"
    gold = "#f4b62f"
    cream = "#fff8e8"
    light_blue = "#d9ecff"
    white = "#ffffff"
    dark = "#18324a"

    img = Image.new("RGB", (W, H), cream)
    draw = ImageDraw.Draw(img)
    title_font = font(FONT_BOLD, 66, FONT_ALT_BOLD)
    h_font = font(FONT_BOLD, 44, FONT_ALT_BOLD)
    body_font = font(FONT_REGULAR, 34, FONT_ALT)
    small_font = font(FONT_REGULAR, 28, FONT_ALT)
    cmd_font = font(FONT_BOLD, 31, FONT_ALT_BOLD)

    draw.rectangle((0, 0, W, 245), fill=navy)
    draw.pieslice((-220, 170, 900, 520), 180, 360, fill=light_blue)
    draw.rectangle((0, 210, W, 245), fill=gold)
    draw.text((70, 58), "คัมภีร์การเจนภาพ", font=title_font, fill=white)
    draw.text((74, 144), "Workflow คำสั่งงานร่วมกันสำหรับป้ายประชาสัมพันธ์", font=body_font, fill=white)

    # Main flow cards
    start_y = 300
    card_w = 700
    card_h = 116
    gap_x = 70
    gap_y = 30
    for idx, (command, desc) in enumerate(COMMANDS):
        col = idx % 2
        row = idx // 2
        x = 80 + col * (card_w + gap_x)
        y = start_y + row * (card_h + gap_y)
        fill = white if idx not in (3, 4, 6) else "#fff3cf"
        rounded(draw, (x, y, x + card_w, y + card_h), 26, fill, outline=navy, width=4)
        draw.ellipse((x + 24, y + 25, x + 88, y + 89), fill=gold, outline=navy, width=3)
        num = str(idx + 1)
        num_bbox = draw.textbbox((0, 0), num, font=cmd_font)
        draw.text((x + 56 - (num_bbox[2] - num_bbox[0]) / 2, y + 35), num, font=cmd_font, fill=navy2)
        draw.text((x + 108, y + 20), command, font=cmd_font, fill=navy2)
        draw_wrapped(draw, desc, (x + 108, y + 60), small_font, dark, card_w - 140, line_gap=2)

    lower_y = 1070
    rounded(draw, (80, lower_y, 1520, lower_y + 320), 34, "#eef7ff", outline=navy, width=4)
    draw.text((120, lower_y + 32), "กติกาเหล็ก", font=h_font, fill=navy2)
    rules = [
        "ไม่ข้าม brief → final render",
        "ต้องมี DRAFT LOW หรือ LAYOUT DRAFT ก่อน",
        "QR/โลโก้ใช้ไฟล์จริงเท่านั้น",
        "ข้อความสำคัญห้ามฝังในภาพ AI",
    ]
    yy = lower_y + 94
    for item in rules:
        draw.ellipse((125, yy + 11, 145, yy + 31), fill=gold)
        draw.text((165, yy), item, font=body_font, fill=dark)
        yy += 52

    y2 = 1440
    rounded(draw, (80, y2, 735, y2 + 310), 34, white, outline=gold, width=5)
    draw.text((120, y2 + 34), "AI ทำอะไร", font=h_font, fill=navy2)
    for i, item in enumerate(["ภาพประกอบ", "ตัวละคร", "ฉากห้องสมุด", "mood/tone"]):
        draw.text((130, y2 + 100 + i * 48), f"✓ {item}", font=body_font, fill=dark)

    rounded(draw, (865, y2, 1520, y2 + 310), 34, white, outline="#d9534f", width=5)
    draw.text((905, y2 + 34), "AI ไม่ทำอะไร", font=h_font, fill="#8a1f17")
    for i, item in enumerate(["QR Code", "โลโก้", "ข้อความไทยยาว", "ข้อมูลนโยบายเอง"]):
        draw.text((915, y2 + 100 + i * 48), f"✕ {item}", font=body_font, fill=dark)

    theme_y = 1805
    rounded(draw, (80, theme_y, 1520, theme_y + 250), 34, navy, outline=gold, width=5)
    draw.text((120, theme_y + 36), "Theme ARC/NPU", font=h_font, fill=white)
    draw.text((120, theme_y + 102), "น้ำเงินเข้ม • ทอง/เหลือง • ขาว/ครีม • ฟ้าอ่อน", font=body_font, fill=white)
    draw.text((120, theme_y + 154), "สดใส เป็นมิตร น่ารัก สุภาพ ใช้ example1.jpg เป็น mood/tone เท่านั้น ไม่ลอก layout", font=small_font, fill=white)
    draw.text((120, theme_y + 198), "personal_akeky.png = มุม support เล็ก ๆ ไม่แย่งความเด่นจาก QR และข้อความหลัก", font=small_font, fill=white)

    draw.rectangle((0, H - 90, W, H), fill=navy2)
    draw.text((80, H - 62), "ระบบจองพื้นที่บริการ Smart Creative Learning Space • สำนักวิทยบริการ มหาวิทยาลัยนครพนม", font=small_font, fill=white)
    img.save(PNG_PATH)


if __name__ == "__main__":
    build_infographic()
    build_docx()
    print(DOCX_PATH)
    print(PNG_PATH)
