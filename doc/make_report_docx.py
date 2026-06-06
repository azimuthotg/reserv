"""Generate report-management-2568.docx — รายงานสรุปผลการดำเนินงาน."""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TEAL   = RGBColor(0x22, 0x97, 0x99)
DARK   = RGBColor(0x1a, 0x1f, 0x2e)
GREEN  = RGBColor(0x06, 0xC7, 0x55)
RED    = RGBColor(0xdc, 0x35, 0x45)
GRAY   = RGBColor(0x71, 0x80, 0x96)
ORANGE = RGBColor(0xe6, 0x7e, 0x22)
FONT   = "TH Sarabun New"
MONO   = "Courier New"
SHOTS  = "/mnt/c/projects/reserv/doc/screenshots"

doc = Document()

# ── Page setup (A4) ───────────────────────────────────────────────────────────
sec = doc.sections[0]
sec.page_width    = Cm(21)
sec.page_height   = Cm(29.7)
sec.top_margin    = Cm(2.5)
sec.bottom_margin = Cm(2.5)
sec.left_margin   = Cm(3.0)
sec.right_margin  = Cm(2.0)

# ── Default style ─────────────────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = FONT
style.font.size = Pt(14)
style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT)

def _set_heading_style(level, size, color=DARK):
    s = doc.styles[f'Heading {level}']
    s.font.name = FONT
    s.font.size = Pt(size)
    s.font.bold = True
    s.font.color.rgb = color
    s.element.rPr.rFonts.set(qn('w:eastAsia'), FONT)

_set_heading_style(1, 20, DARK)
_set_heading_style(2, 17, TEAL)
_set_heading_style(3, 15, DARK)

# ── Helpers ───────────────────────────────────────────────────────────────────
def _run(para_obj, text, bold=False, size=14, color=None, italic=False, font=FONT):
    r = para_obj.add_run(text)
    r.font.name = font
    r.font.bold = bold
    r.font.size = Pt(size)
    r.font.italic = italic
    r.element.rPr.rFonts.set(qn('w:eastAsia'), font)
    if color:
        r.font.color.rgb = color
    return r

def para(text="", bold=False, size=14, color=None,
         align=WD_ALIGN_PARAGRAPH.LEFT,
         space_before=0, space_after=6, italic=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if text:
        _run(p, text, bold=bold, size=size, color=color, italic=italic)
    return p

def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    h.paragraph_format.space_after  = Pt(6)
    for r in h.runs:
        r.font.name = FONT
        r.element.rPr.rFonts.set(qn('w:eastAsia'), FONT)

def bullet(text, indent=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5 + indent * 0.5)
    p.paragraph_format.space_after = Pt(3)
    _run(p, text, size=14)

def page_break():
    doc.add_page_break()

def add_image(filename, width_cm=14, caption=None):
    path = os.path.join(SHOTS, filename)
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(4)
        p.add_run().add_picture(path, width=Cm(width_cm))
    else:
        para(f"[ภาพ: {filename} — ไม่พบไฟล์]", italic=True, color=GRAY)
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(10)
        _run(cp, caption, size=12, italic=True, color=GRAY)

def simple_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    hrow = t.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        _shade_cell(cell, '1A1F2E')
        cp = cell.paragraphs[0]
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(cp, h, bold=True, size=12, color=RGBColor(0xFF, 0xFF, 0xFF))
    for ri, row_data in enumerate(rows):
        row = t.rows[ri + 1]
        fill = 'F0F8F8' if ri % 2 == 0 else 'FFFFFF'
        for ci, cell_text in enumerate(row_data):
            cell = row.cells[ci]
            _shade_cell(cell, fill)
            _run(cell.paragraphs[0], str(cell_text), size=13)
    if col_widths:
        for ci, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[ci].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return t

def _shade_cell(cell, fill_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def info_box(text, fill='E6F4F4', text_color=None):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = t.cell(0, 0)
    _shade_cell(cell, fill)
    cp = cell.paragraphs[0]
    cp.paragraph_format.space_before = Pt(5)
    cp.paragraph_format.space_after  = Pt(5)
    _run(cp, text, size=13, color=text_color or DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def arch_box(lines):
    """Architecture diagram in monospace shaded box."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = t.cell(0, 0)
    _shade_cell(cell, 'F4F4F4')
    cp = cell.paragraphs[0]
    cp.paragraph_format.space_before = Pt(6)
    cp.paragraph_format.space_after  = Pt(6)
    cp.paragraph_format.left_indent  = Cm(0.3)
    first = True
    for line in lines:
        if first:
            _run(cp, line, size=10, font=MONO, color=DARK)
            first = False
        else:
            new_p = OxmlElement('w:p')
            cell._tc.append(new_p)
            from docx.oxml.ns import qn as _qn
            pPr = OxmlElement('w:pPr')
            ind = OxmlElement('w:ind')
            ind.set(_qn('w:left'), '216')
            pPr.append(ind)
            new_p.append(pPr)
            rpr = OxmlElement('w:rPr')
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(_qn('w:ascii'), MONO)
            rFonts.set(_qn('w:hAnsi'), MONO)
            sz = OxmlElement('w:sz')
            sz.set(_qn('w:val'), '20')
            rpr.append(rFonts)
            rpr.append(sz)
            run_el = OxmlElement('w:r')
            run_el.append(rpr)
            t_el = OxmlElement('w:t')
            t_el.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            t_el.text = line
            run_el.append(t_el)
            new_p.append(run_el)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

def caption_label(text):
    """รูปที่ X caption below diagram."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    _run(p, text, size=12, italic=True, color=GRAY)


# ═══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════════════════════════════════
para(space_before=60)

para("รายงานสรุปผลการดำเนินงาน", bold=True, size=28, color=DARK,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0)
para("โครงการพัฒนาระบบควบคุมพื้นที่เรียนรู้อัตโนมัติ", bold=True, size=20,
     color=TEAL, align=WD_ALIGN_PARAGRAPH.CENTER)
para("Smart Creative Learning Space", bold=True, size=18, color=TEAL,
     align=WD_ALIGN_PARAGRAPH.CENTER)

para(space_before=40)

simple_table(
    ["รายการ", "รายละเอียด"],
    [
        ["หน่วยงาน", "สำนักวิทยบริการ มหาวิทยาลัยนครพนม"],
        ["ปีงบประมาณ", "2568"],
        ["ผู้จัดทำ", "………………………………………………"],
        ["ผู้รับผิดชอบโครงการ", "………………………………………………"],
        ["วันที่จัดทำ", "พฤษภาคม 2568"],
    ],
    col_widths=[5, 10]
)

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# บทสรุปผู้บริหาร
# ═══════════════════════════════════════════════════════════════════════════════
heading("บทสรุปผู้บริหาร (Executive Summary)", 1)

para("สำนักวิทยบริการ มหาวิทยาลัยนครพนม ได้ดำเนินโครงการพัฒนาระบบจองพื้นที่บริการ Smart Creative Learning Space ภายใต้กรอบแนวคิด Design Thinking โดยมีจุดมุ่งหมายหลักเพื่อแก้ไขข้อจำกัดของระบบเดิมที่พึ่งพาบริการคลาวด์ภายนอก (Google Apps Script และ Google Sheets) ซึ่งส่งผลให้ระบบมีความล่าช้า ขาดความยืดหยุ่น และมีความเสี่ยงด้านความต่อเนื่องของการให้บริการ")

para("โครงการแบ่งออกเป็น 2 กิจกรรมหลัก กิจกรรมที่ 1 มุ่งพัฒนาระบบจองใหม่บนโครงสร้างพื้นฐานของมหาวิทยาลัย (On-premise) ด้วย Django Framework และฐานข้อมูล MySQL โดยยังรักษาการเชื่อมต่อกับ LINE LIFF เพื่อประสบการณ์ผู้ใช้ที่คุ้นเคย พร้อมเพิ่มช่องทางการจองผ่านเว็บไซต์สำนักฯ กิจกรรมที่ 2 เชื่อมต่อระบบจองกับอุปกรณ์ IoT (Sonoff Smart Plug) ผ่าน Home Assistant เพื่อแจ้งเตือนสถานะอุปกรณ์และส่งการแจ้งเตือนให้ผู้จองผ่าน LINE Messaging API อัตโนมัติ")

para("ปัจจุบันทั้ง 2 กิจกรรมดำเนินการเสร็จสมบูรณ์และเปิดใช้งานจริงที่ https://lib.npu.ac.th/reserv/ แล้ว รองรับผู้ใช้ทั้งกลุ่มนักศึกษาและบุคลากรผ่าน 2 ช่องทาง (LINE OA และเว็บไซต์) โดยไม่มีค่าใช้จ่ายรายเดือนสำหรับบริการคลาวด์ภายนอกอีกต่อไป")

para("สรุปสถานะกิจกรรม:", bold=True, space_before=6)
simple_table(
    ["กิจกรรม", "ชื่อ", "สถานะ"],
    [
        ["กิจกรรมที่ 1", "ปรับปรุงฐานข้อมูลและระบบจอง (Local Booking Core)", "✅ เสร็จสมบูรณ์ — ใช้งานจริงแล้ว"],
        ["กิจกรรมที่ 2", "ระบบแจ้งเตือนและตรวจสอบอุปกรณ์ (Alert-as-a-Service)", "✅ เสร็จสมบูรณ์ — ใช้งานจริงแล้ว"],
    ],
    col_widths=[3, 8.5, 5.5]
)

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 1. ความเป็นมาและปัญหา
# ═══════════════════════════════════════════════════════════════════════════════
heading("1. ความเป็นมาและปัญหา", 1)

heading("1.1 บริบทการให้บริการ", 2)
para("สำนักวิทยบริการ มหาวิทยาลัยนครพนม จัดพื้นที่บริการ Smart Creative Learning Space เพื่อให้นักศึกษาและบุคลากรสามารถใช้พื้นที่สำหรับการศึกษา ทำงานกลุ่ม และกิจกรรมเชิงสร้างสรรค์ โดยพื้นที่บริการประกอบด้วยห้องทั้งหมด 5 ประเภท ดังนี้")
simple_table(
    ["ชื่อห้อง", "ลักษณะการใช้งาน"],
    [
        ["Mini Theater", "ห้องฉายภาพยนตร์ขนาดเล็ก เหมาะสำหรับการนำเสนอและกิจกรรมกลุ่ม"],
        ["Edutainment Zone", "พื้นที่สื่อการเรียนรู้เชิงบันเทิง"],
        ["Canva Studio", "ห้องออกแบบกราฟิกและสื่อดิจิทัล"],
        ["AI Learning Zone (Chat-GPT Lab)", "ห้องทดลองใช้งาน AI และเทคโนโลยีปัญญาประดิษฐ์"],
        ["Meeting Room F1", "ห้องประชุมชั้น 1"],
    ],
    col_widths=[5, 11]
)

heading("1.2 ระบบเดิมและข้อจำกัด", 2)
para("ระบบจองพื้นที่เดิมถูกพัฒนาในลักษณะ Quick Solution โดยใช้เครื่องมือที่สามารถเริ่มต้นได้ทันทีโดยไม่ต้องใช้งบประมาณเพิ่มเติม สถาปัตยกรรมของระบบเดิมมีลักษณะดังแผนภาพ:")

caption_label("รูปที่ 1 — สถาปัตยกรรมระบบเดิม (ก่อนพัฒนา)")
arch_box([
    "ผู้ใช้ใน LINE OA",
    "    ↓ กด Flex Message Card → เปิด LIFF",
    "HTML Form  (arc.npu.ac.th — external hosting)",
    "    ↓ ผู้ใช้กรอกข้อมูลด้วยตนเอง",
    "Google Apps Script (Web App — Google Cloud)",
    "    ↓ บันทึกข้อมูล",
    "Google Sheets (ฐานข้อมูล — Google Cloud)",
    "    ↓ ตรวจสิทธิ์ผ่าน",
    "testdb.php (external server)",
    "    ↓ ดึงข้อมูลผู้ใช้จาก",
    "ฐานข้อมูล LDAP / Active Directory ของมหาวิทยาลัย",
])

para("ระบบเดิมพบข้อจำกัดสำคัญ 4 ด้าน:", bold=True, space_before=4)

bullet("การพึ่งพาบริการภายนอก (Vendor Lock-in): Google Sheets มีข้อจำกัดด้านจำนวน API calls ต่อวัน, Google Apps Script มี execution time limit (สูงสุด 6 นาที/ครั้ง), และ LINE Notify ปิดบริการถาวรตั้งแต่เมษายน 2568 ทำให้ระบบแจ้งเตือนหยุดทำงานโดยสมบูรณ์")
bullet("ความปลอดภัยของข้อมูล: ข้อมูลผู้ใช้และการจองถูกจัดเก็บบน Google Cloud ภายนอกองค์กร ไม่สอดคล้องกับนโยบายความปลอดภัยขององค์กร และขาดกลไก audit trail ที่เป็นมาตรฐาน")
bullet("ข้อจำกัดด้านฟังก์ชัน: ไม่สามารถตรวจสอบ conflict การจองพร้อมกันได้อย่างน่าเชื่อถือ, ผู้ใช้ต้องกรอกชื่อ-คณะ-สาขาด้วยตนเองทุกครั้ง, ไม่มีระบบจัดการวันหยุด และไม่มีปฏิทินรวม")
bullet("ความยากในการบำรุงรักษา: โค้ดไม่มี version control ทำให้ไม่สามารถ rollback ได้, การ debug ทำได้ยากเนื่องจากไม่มี structured log")

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 2. กรอบแนวคิดการพัฒนา
# ═══════════════════════════════════════════════════════════════════════════════
heading("2. กรอบแนวคิดการพัฒนา — Design Thinking", 1)
para("โครงการนี้นำกระบวนการคิดเชิงออกแบบ (Design Thinking) 5 ขั้นตอนมาเป็นกรอบในการพัฒนา เพื่อให้ระบบใหม่ตอบสนองความต้องการของผู้ใช้จริงได้อย่างตรงจุด")
simple_table(
    ["ขั้นตอน", "กิจกรรม", "ผลลัพธ์"],
    [
        ["1. Empathize", "สัมภาษณ์เจ้าหน้าที่และผู้ใช้บริการ, วิเคราะห์ปัญหาที่พบบ่อย", "รายการปัญหาจากมุมมองผู้ใช้จริง"],
        ["2. Define", "จำกัดขอบเขตปัญหาหลัก: ความล่าช้า, ข้อมูลอยู่นอกองค์กร, ขาดการแจ้งเตือน", "Problem Statement ที่ชัดเจน"],
        ["3. Ideate", "ออกแบบสถาปัตยกรรมใหม่, เลือก technology stack, วางแผนกิจกรรม", "แผนงานและสถาปัตยกรรมระบบ"],
        ["4. Prototype", "พัฒนาระบบตาม Phase A–E, ทดสอบแต่ละส่วน", "ระบบ prototype บน staging"],
        ["5. Test", "ทดสอบกับผู้ใช้จริง, ปรับแก้ตามความเห็น, deploy production", "ระบบ production ที่ใช้งานได้จริง"],
    ],
    col_widths=[3, 7, 7]
)
para("ผลจากกระบวนการ Design Thinking นำไปสู่การกำหนด 2 กิจกรรมหลัก:")
bullet("กิจกรรมที่ 1 — Local Booking Core: ปรับปรุงฐานข้อมูลและระบบจองให้อยู่ภายในองค์กร")
bullet("กิจกรรมที่ 2 — Alert-as-a-Service: ระบบแจ้งเตือนและตรวจสอบสถานะอุปกรณ์อัตโนมัติ")

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 3. กิจกรรมที่ 1
# ═══════════════════════════════════════════════════════════════════════════════
heading("3. กิจกรรมที่ 1 — ปรับปรุงฐานข้อมูลและระบบจอง (Local Booking Core)", 1)

heading("3.1 ความต้องการระบบ", 2)
para("จากการวิเคราะห์ระบบเดิมและรับฟังความต้องการของผู้ใช้ กำหนดความต้องการระบบดังนี้", space_after=4)
para("ความต้องการเชิงฟังก์ชัน (Functional Requirements):", bold=True, space_after=2)
simple_table(
    ["รหัส", "ความต้องการ", "ระดับ"],
    [
        ["FR-01", "ผู้ใช้จองห้องผ่าน LINE LIFF ได้โดยไม่ต้องออกจาก LINE", "สูงสุด"],
        ["FR-02", "ตรวจสอบสิทธิ์ผ่าน LINE userId เชื่อมกับบัญชีมหาวิทยาลัย", "สูงสุด"],
        ["FR-03", "ดึงข้อมูลชื่อ-นามสกุล คณะ สาขา จากระบบมหาวิทยาลัยอัตโนมัติ", "สูง"],
        ["FR-04", "ตรวจสอบ conflict การจองแบบ real-time ป้องกันการจองซ้ำซ้อน", "สูงสุด"],
        ["FR-05", "บริหารจัดการวันหยุดและการปิดบริการชั่วคราวผ่าน Admin Panel", "สูง"],
        ["FR-06", "จำกัดการจองล่วงหน้าไม่เกิน 3 วันทำงาน", "สูง"],
        ["FR-07", "แจ้งเตือนจองสำเร็จผ่าน LINE ทันทีหลังจองเสร็จ", "สูง"],
        ["FR-08", "แจ้งเตือนผู้จองก่อนถึงเวลาใช้งาน 15 นาที", "ปานกลาง"],
        ["FR-09", "แจ้งเตือนก่อนหมดเวลาใช้งาน 10 นาที", "ปานกลาง"],
        ["FR-10", "ผู้ดูแลระบบจัดการข้อมูลการจองและผู้ใช้ผ่าน Admin Panel", "สูง"],
    ],
    col_widths=[2, 11.5, 2.5]
)
para("ความต้องการเชิงไม่ใช่ฟังก์ชัน (Non-Functional Requirements):", bold=True, space_after=2)
simple_table(
    ["รหัส", "ความต้องการ", "เป้าหมาย"],
    [
        ["NFR-01", "ข้อมูลทั้งหมดต้องจัดเก็บบนเซิร์ฟเวอร์ภายในองค์กร", "100% on-premise"],
        ["NFR-02", "รองรับผู้ใช้งานพร้อมกันได้เพียงพอสำหรับองค์กร", "≥ 50 concurrent users"],
        ["NFR-03", "เวลาตอบสนองของระบบ", "< 2 วินาที"],
        ["NFR-04", "ความพร้อมใช้งาน", "≥ 95% uptime"],
        ["NFR-05", "รองรับ LINE In-App Browser (mobile) และ External Browser (desktop)", "ครบทั้ง 2 แพลตฟอร์ม"],
    ],
    col_widths=[2, 10, 5]
)

heading("3.2 สถาปัตยกรรมระบบใหม่", 2)
para("ระบบใหม่ย้ายฐานข้อมูลและ logic ทั้งหมดมาสู่โครงสร้างพื้นฐานภายในมหาวิทยาลัย ดังแผนภาพ:")

caption_label("รูปที่ 2 — สถาปัตยกรรมระบบใหม่ (หลังพัฒนา)")
arch_box([
    "ช่องทางที่ 1: ผู้ใช้ใน LINE OA (มือถือ)      ช่องทางที่ 2: ผู้ใช้ทั่วไป (Desktop/Browser)",
    "    ↓ กด 'จองห้อง' ใน Rich Menu                   ↓ เปิด https://lib.npu.ac.th/reserv/",
    "    ↓ URL: https://lib.npu.ac.th/reserv/booking/?room=X",
    "",
    "                IIS ARR — lib.npu.ac.th (Reverse Proxy)",
    "                    ↓ /reserv/* → 127.0.0.1:8003",
    "",
    "                Waitress WSGI Server (port 8003)",
    "                ├── /reserv/static/* → WhiteNoise",
    "                └── /reserv/*        → Django Application",
    "",
    "                Django Application (booking/views.py)",
    "                ├── booking_page()    → แสดงฟอร์มจอง",
    "                ├── check_user()     → ตรวจสิทธิ์ + ดึงโปรไฟล์",
    "                ├── create_booking() → ตรวจ conflict + บันทึก",
    "                └── admin/           → Django Admin Panel",
    "",
    "    External APIs              MySQL Database           Windows Services",
    "    (api.npu.ac.th)            (On-premise)             - NSSM (Waitress)",
    "    - /api/{userId}/           - booking_room           - Task Scheduler",
    "    - /std-info/{ldap}/        - booking_lineuser         (send_reminders)",
    "    - /staff-info/{ldap}/      - booking_booking",
    "    - api.line.me (push)       - booking_bookinglog",
    "                               - booking_holidaydate",
])

para("เปรียบเทียบสถาปัตยกรรมระบบเดิมและระบบใหม่:", bold=True, space_before=6, space_after=2)
simple_table(
    ["องค์ประกอบ", "ระบบเดิม", "ระบบใหม่"],
    [
        ["Web Application", "HTML Form บน external hosting", "Django App บน Windows Server ภายใน"],
        ["ฐานข้อมูล", "Google Sheets (Google Cloud)", "MySQL 8.0 (on-premise)"],
        ["ตรวจสิทธิ์", "testdb.php (external server)", "api.npu.ac.th (university API)"],
        ["การแจ้งเตือน", "LINE Notify (ปิดบริการแล้ว)", "LINE Messaging API Push Message"],
        ["Process Manager", "ไม่มี", "NSSM Windows Service"],
        ["Version Control", "ไม่มี", "Git (GitHub)"],
    ],
    col_widths=[4, 6, 7]
)

heading("3.3 Technology Stack และเหตุผลการเลือก", 2)
simple_table(
    ["ส่วนประกอบ", "ที่เลือกใช้", "เหตุผล"],
    [
        ["Backend Framework", "Django 4.2 / Python 3.12", "มี ORM, Admin Panel, Migration ในตัว เหมาะสำหรับทีมขนาดเล็ก"],
        ["Database", "MySQL 8.0", "ใช้ infrastructure ที่มีอยู่แล้วภายในองค์กร"],
        ["WSGI Server", "Waitress", "รองรับ Windows Server โดยตรง ไม่ต้องพึ่ง Linux"],
        ["Reverse Proxy", "IIS ARR", "Windows Server มี IIS อยู่แล้ว ไม่ต้องติดตั้งเพิ่ม"],
        ["Frontend", "Bootstrap 5.3 + FullCalendar v6", "ลด complexity เหมาะกับ LIFF single-page"],
        ["LINE Integration", "LIFF SDK 2.15", "ดึง userId โดยตรงภายใน LINE app"],
        ["Process Manager", "NSSM", "จัดการ Python process เป็น Windows Service ได้ง่าย"],
    ],
    col_widths=[4, 4.5, 8.5]
)

heading("3.4 ช่องทางการจอง 2 ช่องทาง", 2)
para("ช่องทางที่ 1 — LINE Official Account (มือถือ):", bold=True, space_after=2)
bullet("ผู้ใช้กดปุ่ม 'จองห้อง' ใน Rich Menu ของ LINE OA สำนักวิทยบริการ")
bullet("ระบบเปิดหน้าจองผ่าน LINE LIFF ดึง userId จาก LINE โดยอัตโนมัติ")
bullet("หากยังไม่ผูกบัญชี ระบบนำไปยังหน้าลงทะเบียน (ทำครั้งเดียว)")
bullet("ระบบดึงชื่อ-คณะ-สาขาจากระบบมหาวิทยาลัยอัตโนมัติ ผู้ใช้เพียงเลือกวัน/เวลาและยืนยัน")
para("ช่องทางที่ 2 — เว็บไซต์สำนักวิทยบริการ (Desktop/Browser):", bold=True, space_before=4, space_after=2)
bullet("ผู้ใช้เปิดเว็บไซต์ https://lib.npu.ac.th/reserv/ ผ่าน browser")
bullet("ล็อกอินด้วยบัญชี LDAP ของมหาวิทยาลัย (username/password เดิม)")
bullet("ระบบแสดงฟอร์มจองพร้อมปฏิทินการจองทั้งหมด")

heading("3.5 การออกแบบฐานข้อมูล", 2)
simple_table(
    ["ตาราง", "บทบาท", "ข้อมูลสำคัญ"],
    [
        ["booking_room", "ข้อมูลห้องบริการ", "ชื่อห้อง, เวลาเปิด-ปิด, สถานะ, รหัส IoT"],
        ["booking_lineuser", "Cache ผู้ใช้ที่ผูก LINE กับ LDAP", "LINE userId, ชื่อจริง, คณะ, สาขา, เวลา cache"],
        ["booking_booking", "บันทึกการจองทั้งหมด", "ห้อง, ผู้จอง, วันที่, เวลาเริ่ม-สิ้นสุด, สถานะ, tracking แจ้งเตือน"],
        ["booking_bookinglog", "Audit Trail ทุก action", "การจอง, action, หมายเหตุ, timestamp"],
        ["booking_holidaydate", "วันหยุดและปิดบริการ", "วันที่, คำอธิบาย, สถานะ"],
    ],
    col_widths=[4, 5, 8]
)
info_box("กลไก Conflict Check: ระบบใช้ SELECT FOR UPDATE ร่วมกับ transaction.atomic() เพื่อป้องกัน race condition กรณีผู้ใช้หลายคนจองช่วงเวลาเดียวกันพร้อมกัน — จุดอ่อนสำคัญที่ระบบเดิมไม่สามารถป้องกันได้")
info_box("Profile Cache: ระบบบันทึกข้อมูลชื่อ-คณะ-สาขาไว้ใน database และ refresh ทุก 30 วัน ลดเวลาโหลดจาก ~800ms เหลือ ~50ms สำหรับผู้ใช้เดิม")

page_break()

# ── Screenshots กิจกรรม 1 ────────────────────────────────────────────────────
heading("3.6 หน้าจอระบบ", 2)

para("รูปที่ 3 — หน้าลงทะเบียนผูกบัญชี LINE กับมหาวิทยาลัย (มือถือ — LIFF)", bold=True, space_after=2)
add_image("01_register.png", width_cm=8,
          caption="ผู้ใช้ที่ยังไม่เคยลงทะเบียนจะถูกนำมายังหน้านี้ก่อน เพื่อผูก LINE Account กับ username/password ของมหาวิทยาลัย (ทำครั้งเดียว ระบบจดจำอัตโนมัติ)")

para("รูปที่ 4 — หน้าแรกของระบบจอง", bold=True, space_before=8, space_after=2)
add_image("02_landing.png", width_cm=14,
          caption="หน้าแรกแสดงภาพรวมห้องบริการและสถานะการใช้งาน ผู้ใช้สามารถเลือกห้องที่ต้องการจองได้จากหน้านี้")

para("รูปที่ 5 — ฟอร์มจองพื้นที่ (มือถือ — ผ่าน LINE LIFF)", bold=True, space_before=8, space_after=2)
add_image("03_booking.png", width_cm=8,
          caption="ฟอร์มจองแสดงผลใน LINE app บนมือถือ ระบบดึงชื่อ-คณะ-สาขาอัตโนมัติ ผู้ใช้เพียงเลือกวันที่ เวลา และจำนวนผู้เข้าใช้")

para("รูปที่ 6 — หน้ายืนยันการจองสำเร็จ (มือถือ — LIFF)", bold=True, space_before=8, space_after=2)
add_image("04_success.png", width_cm=8,
          caption="เมื่อจองสำเร็จ ระบบแสดงสรุปรายละเอียดและส่ง LINE notification ให้ผู้จองทันที")

page_break()

para("รูปที่ 7 — ปฏิทินการจอง (Desktop — เว็บไซต์)", bold=True, space_after=2)
add_image("06_calendar.png", width_cm=15,
          caption="ปฏิทินแสดงการจองทั้งหมดรูปแบบรายวัน รายสัปดาห์ และรายเดือน รองรับ desktop browser ผู้ใช้และเจ้าหน้าที่ตรวจสอบความว่างก่อนจองได้")

para("รูปที่ 8 — Admin Dashboard — ภาพรวมระบบ (Desktop)", bold=True, space_before=8, space_after=2)
add_image("admin_02_dashboard.png", width_cm=15,
          caption="หน้า Admin Panel สำหรับเจ้าหน้าที่ดูแลระบบ แสดงสถิติการจอง จำนวนห้อง และการใช้งานระบบ")

para("รูปที่ 9 — หน้าจัดการรายการจอง — Admin (Desktop)", bold=True, space_before=8, space_after=2)
add_image("admin_03_bookings.png", width_cm=15,
          caption="เจ้าหน้าที่สามารถดู แก้ไข และยกเลิกการจองได้จากหน้านี้ พร้อม filter ตามเงื่อนไขต่าง ๆ")

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 4. กิจกรรมที่ 2
# ═══════════════════════════════════════════════════════════════════════════════
heading("4. กิจกรรมที่ 2 — ระบบแจ้งเตือนและตรวจสอบอุปกรณ์ (Alert-as-a-Service)", 1)

heading("4.1 ที่มาและความจำเป็น", 2)
para("LINE Notify ปิดบริการถาวรตั้งแต่วันที่ 1 เมษายน 2568 ทำให้ระบบแจ้งเตือนเดิมหยุดทำงานโดยสมบูรณ์ กิจกรรมนี้จึงพัฒนาระบบแจ้งเตือนใหม่ผ่าน LINE Messaging API (Push Message) ซึ่งเป็นบริการระดับ Official ของ LINE ที่มีความเสถียรสูง พร้อมกันนี้ยังเชื่อมต่อกับอุปกรณ์ IoT ในพื้นที่บริการเพื่อการจัดการอัตโนมัติ")

heading("4.2 ระบบแจ้งเตือน LINE Messaging API", 2)
para("ระบบส่งการแจ้งเตือนผ่าน LINE อัตโนมัติทั้งหมด 7 กรณี:")
simple_table(
    ["#", "กรณี", "เวลาส่ง", "รูปแบบ / เนื้อหา"],
    [
        ["1", "จองสำเร็จ", "ทันทีหลังจองเสร็จ", "Flex Message + ปุ่ม 'ดูการจองของฉัน'"],
        ["2", "แจ้งเตือนก่อนเริ่ม 15 นาที", "ก่อนเวลาจอง 15 นาที", "Flex Message + ปุ่ม 'Check-in เลย' พร้อมแจ้งกำหนดเส้นตาย check-in"],
        ["3", "Check-in สำเร็จ", "ทันทีที่ผู้จอง check-in", "Flex Message + ปุ่ม 'ควบคุมอุปกรณ์ในห้อง'"],
        ["4", "ยกเลิกอัตโนมัติ", "15 นาทีหลังเวลาเริ่ม (ถ้าไม่ check-in)", "Text — แจ้งว่าการจองถูกยกเลิกเนื่องจากไม่ check-in"],
        ["5", "แจ้งเตือนก่อนหมด 10 นาที", "ก่อนหมดเวลา 10 นาที", "Text — เตือนให้เตรียมเก็บของและออกจากห้อง"],
        ["6", "ยกเลิกโดยผู้ใช้เอง", "ทันทีที่กดยกเลิก", "Text — สรุปรายละเอียดการจองที่ยกเลิก"],
        ["7", "ยกเลิกโดยเจ้าหน้าที่", "ทันทีที่เจ้าหน้าที่ยกเลิก", "Text — แจ้งเหตุผลการยกเลิก (ถ้ามี) + ช่องทางติดต่อ"],
    ],
    col_widths=[0.8, 4.2, 4, 8]
)
info_box("กรณีที่ 2, 4, 5 ทำงานผ่าน Windows Task Scheduler รันทุก 1 นาที มีกลไกป้องกันการแจ้งเตือนซ้ำด้วย flag (notified_15min, notified_10min, notified_auto_off) ในฐานข้อมูล")

heading("4.3 การเชื่อมต่ออุปกรณ์ IoT (Sonoff Integration)", 2)
para("ระบบเชื่อมต่อกับ Smart Plug ยี่ห้อ Sonoff ที่ติดตั้งในแต่ละห้องบริการ ผ่านแพลตฟอร์ม Home Assistant")

caption_label("รูปที่ 2ข — การทำงานของ IoT Integration")
arch_box([
    "ระบบจอง (Django)",
    "    ↓ booking สถานะ confirmed",
    "/api/check-access/ endpoint (protected ด้วย X-HA-Token)",
    "    ↓",
    "Home Assistant (HA)",
    "    ↓ trigger automation",
    "Sonoff Smart Plug",
    "    ↓ เปิด/ปิดอุปกรณ์ในห้อง",
    "LINE notification → ผู้ดูแลระบบ (แจ้งเตือนสถานะอัปเดต)",
])

bullet("Endpoint ที่ใช้: /api/check-access/ ป้องกันด้วย token (X-HA-Token) เพื่อความปลอดภัย")
bullet("การเชื่อมต่อ: ผ่าน api.npu.ac.th/sonoff/ เป็น API กลางของมหาวิทยาลัย")
bullet("การตั้งค่า: แต่ละห้องมีรหัส ha_entity_id ในฐานข้อมูล Room สามารถเพิ่มห้องใหม่ได้โดยไม่ต้องแก้โค้ด")

heading("4.4 ตัวอย่างการแจ้งเตือน", 2)
para("รูปที่ 10 — ตัวอย่างการแจ้งเตือนผ่าน LINE", bold=True, space_after=2)
add_image("08_notification.png", width_cm=8,
          caption="ตัวอย่าง LINE notification ที่ผู้จองได้รับ แสดงรายละเอียดการจองครบถ้วน ส่งผ่าน LINE Messaging API ซึ่งมีความเสถียรกว่า LINE Notify ที่ปิดบริการแล้ว")

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 5. เปรียบเทียบระบบ
# ═══════════════════════════════════════════════════════════════════════════════
heading("5. เปรียบเทียบระบบเดิมและระบบใหม่", 1)
simple_table(
    ["หัวข้อเปรียบเทียบ", "ระบบเดิม", "ระบบใหม่"],
    [
        ["การเก็บข้อมูล", "Google Sheets (cloud ภายนอก)", "MySQL บน server ภายในองค์กร"],
        ["ตำแหน่งข้อมูล", "ภายนอกองค์กร (Google)", "ภายในองค์กร (on-premise)"],
        ["Conflict Check", "ไม่น่าเชื่อถือ (race condition)", "SELECT FOR UPDATE + transaction.atomic()"],
        ["ตรวจสิทธิ์", "testdb.php (external server)", "api.npu.ac.th + profile cache 30 วัน"],
        ["การดึงข้อมูลผู้ใช้", "กรอกเองทุกครั้ง", "ดึงจากระบบมหาวิทยาลัยอัตโนมัติ"],
        ["การจัดการวันหยุด", "ไม่มี", "HolidayDate model + Admin UI"],
        ["ข้อจำกัดการจองล่วงหน้า", "ไม่มี", "3 วันทำงาน (คำนวณอัตโนมัติ)"],
        ["แจ้งเตือนจองสำเร็จ", "ไม่มี", "LINE push ทันที (Flex Message)"],
        ["แจ้งเตือนก่อนเริ่ม 15 นาที + ปุ่ม Check-in", "ไม่มี", "LINE push (Task Scheduler, Flex Message)"],
        ["Check-in สำเร็จ + ปุ่มควบคุมอุปกรณ์", "ไม่มี", "LINE push ทันที (Flex Message)"],
        ["ยกเลิกอัตโนมัติ (ไม่ check-in)", "ไม่มี", "LINE push อัตโนมัติ (Task Scheduler)"],
        ["แจ้งเตือนก่อนหมด 10 นาที", "ไม่มี", "LINE push (Task Scheduler)"],
        ["ยกเลิกโดยผู้ใช้ / เจ้าหน้าที่", "ไม่มี", "LINE push ทันที พร้อมเหตุผล"],
        ["IoT Integration", "ไม่มี", "Sonoff via Home Assistant"],
        ["Audit Trail", "ไม่มี", "BookingLog ทุก action"],
        ["Admin Panel", "Google Sheets (จำกัด)", "Django Admin (เต็มรูปแบบ)"],
        ["Version Control", "ไม่มี", "Git / GitHub"],
        ["ค่าบริการรายเดือน", "$0 แต่มีข้อจำกัดและความเสี่ยง", "$0 (on-premise ไม่มีค่าบริการ cloud)"],
        ["ช่องทางการจอง", "LINE LIFF เท่านั้น", "LINE LIFF + เว็บไซต์ (2 ช่องทาง)"],
    ],
    col_widths=[5.5, 5.5, 6]
)

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 6. ผลการดำเนินงาน
# ═══════════════════════════════════════════════════════════════════════════════
heading("6. ผลการดำเนินงานและสถานะฟีเจอร์", 1)

heading("6.1 สถานะกิจกรรมหลัก", 2)
simple_table(
    ["กิจกรรม", "สถานะ", "URL"],
    [
        ["กิจกรรมที่ 1 — Local Booking Core", "✅ เสร็จสมบูรณ์ — ใช้งานจริงแล้ว", "https://lib.npu.ac.th/reserv/"],
        ["กิจกรรมที่ 2 — Alert-as-a-Service", "✅ เสร็จสมบูรณ์ — ใช้งานจริงแล้ว", "—"],
    ],
    col_widths=[6, 5.5, 5.5]
)

heading("6.2 สถานะฟีเจอร์ทั้งหมด", 2)
simple_table(
    ["ฟีเจอร์", "สถานะ"],
    [
        ["จองห้อง + conflict check แบบ real-time", "✅ พร้อมใช้งาน"],
        ["ตรวจสิทธิ์ผ่าน NPU API + profile cache", "✅ พร้อมใช้งาน"],
        ["รองรับนักศึกษาและบุคลากรในระบบเดียว", "✅ พร้อมใช้งาน"],
        ["วันหยุดราชการ + ข้อจำกัด 3 วันทำงาน", "✅ พร้อมใช้งาน"],
        ["แจ้งเตือนจองสำเร็จผ่าน LINE ทันที (Flex Message)", "✅ พร้อมใช้งาน"],
        ["แจ้งเตือนก่อนเริ่ม 15 นาที + ปุ่ม Check-in", "✅ พร้อมใช้งาน"],
        ["Check-in สำเร็จ + ปุ่มควบคุมอุปกรณ์", "✅ พร้อมใช้งาน"],
        ["ยกเลิกอัตโนมัติเมื่อไม่ check-in ใน 15 นาที", "✅ พร้อมใช้งาน"],
        ["แจ้งเตือนก่อนหมดเวลา 10 นาที", "✅ พร้อมใช้งาน"],
        ["แจ้งเตือนยกเลิกการจอง (ผู้ใช้ / เจ้าหน้าที่)", "✅ พร้อมใช้งาน"],
        ["IoT Sonoff integration via Home Assistant", "✅ พร้อมใช้งาน"],
        ["ปฏิทินรวม (FullCalendar)", "✅ พร้อมใช้งาน"],
        ["Admin Panel สำหรับเจ้าหน้าที่", "✅ พร้อมใช้งาน"],
        ["Audit Trail (BookingLog)", "✅ พร้อมใช้งาน"],
        ["การจองผ่านเว็บไซต์ (ไม่ผ่าน LINE)", "✅ พร้อมใช้งาน"],
    ],
    col_widths=[13, 4]
)

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 7. Roadmap
# ═══════════════════════════════════════════════════════════════════════════════
heading("7. แผนการพัฒนาต่อเนื่อง (Roadmap)", 1)
para("ภายหลังจากกิจกรรมที่ 1 และ 2 สำเร็จแล้ว ได้วางแผนการพัฒนาเพิ่มเติมในอนาคต")

heading("กิจกรรมที่ 3 — เสริมความพร้อมด้าน Resilience", 2)
bullet("Virtual Card: บัตรเข้าใช้บริการแบบดิจิทัล (JsBarcode) ผู้ใช้แสดงบนมือถือแทนบัตรกระดาษ")
bullet("Walai Integration: เชื่อมต่อกับระบบ Walai เพื่อตรวจสอบสิทธิ์การเข้าห้องสมุดก่อนอนุมัติการจอง")
bullet("UI ยกเลิกการจอง: ปัจจุบันผู้ใช้ยกเลิกได้เฉพาะผ่าน Admin Panel เท่านั้น")
bullet("ระบบ Blacklist: นับจำนวนผู้จองที่ไม่มาใช้บริการ (no-show) และระงับสิทธิ์ชั่วคราว")

heading("กิจกรรมที่ 4 — ขยายพื้นที่และอุปกรณ์ IoT", 2)
bullet("ขยายการเชื่อมต่อ IoT ไปยังห้องบริการเพิ่มเติม")
bullet("ติดตั้งอุปกรณ์ Smart Plug เพิ่มเติมตามงบประมาณที่ได้รับ (40,000 บาท)")
bullet("Dashboard สถิติการใช้งานสำหรับผู้บริหาร")

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 8. สรุป
# ═══════════════════════════════════════════════════════════════════════════════
heading("8. สรุป", 1)
para("โครงการพัฒนาระบบควบคุมพื้นที่เรียนรู้อัตโนมัติ Smart Creative Learning Space ในปีงบประมาณ 2568 ประสบความสำเร็จตามเป้าหมายที่กำหนดทั้ง 2 กิจกรรม")
para("กิจกรรมที่ 1 ดำเนินการย้ายฐานข้อมูลและ logic การทำงานทั้งหมดมาสู่โครงสร้างพื้นฐานภายในมหาวิทยาลัยสำเร็จสมบูรณ์ ระบบใหม่มีความเสถียร ตอบสนองเร็วกว่าระบบเดิม ข้อมูลผู้ใช้และการจองทั้งหมดจัดเก็บภายในองค์กร สอดคล้องกับนโยบายความปลอดภัยของข้อมูล และสามารถขยายฟังก์ชันในอนาคตได้อย่างยืดหยุ่น")
para("กิจกรรมที่ 2 ดำเนินการเชื่อมต่อระบบจองกับอุปกรณ์ IoT และพัฒนาระบบแจ้งเตือนผ่าน LINE Messaging API แทน LINE Notify ที่ปิดบริการแล้ว ช่วยให้ผู้ใช้ได้รับการแจ้งเตือนอัตโนมัติตลอดกระบวนการจอง และเจ้าหน้าที่สามารถติดตามสถานะอุปกรณ์ในพื้นที่บริการได้แบบ real-time")
para("ทั้งสองกิจกรรมเปิดใช้งานจริงที่ https://lib.npu.ac.th/reserv/ รองรับผู้ใช้ทั้งกลุ่มนักศึกษาและบุคลากรผ่าน 2 ช่องทาง (LINE OA และเว็บไซต์) โดยไม่มีค่าใช้จ่ายบริการ cloud ภายนอก")

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# ภาคผนวก
# ═══════════════════════════════════════════════════════════════════════════════
heading("ภาคผนวก — รายการภาพประกอบในรายงาน", 1)
simple_table(
    ["รูปที่", "คำบรรยาย", "ประเภทหน้าจอ"],
    [
        ["รูปที่ 1", "สถาปัตยกรรมระบบเดิม (Diagram)", "—"],
        ["รูปที่ 2", "สถาปัตยกรรมระบบใหม่ (Diagram)", "—"],
        ["รูปที่ 3", "หน้าลงทะเบียนผูกบัญชี LINE", "มือถือ (LIFF)"],
        ["รูปที่ 4", "หน้าแรกของระบบจอง", "—"],
        ["รูปที่ 5", "ฟอร์มจองพื้นที่", "มือถือ (LIFF)"],
        ["รูปที่ 6", "หน้ายืนยันการจองสำเร็จ", "มือถือ (LIFF)"],
        ["รูปที่ 7", "ปฏิทินการจอง", "Desktop (เว็บไซต์)"],
        ["รูปที่ 8", "Admin Dashboard", "Desktop (เว็บไซต์)"],
        ["รูปที่ 9", "รายการจองใน Admin", "Desktop (เว็บไซต์)"],
        ["รูปที่ 10", "ตัวอย่างการแจ้งเตือนผ่าน LINE", "มือถือ (LINE)"],
    ],
    col_widths=[2.5, 9, 5.5]
)

info_box("หมายเหตุ: รูปที่ 3, 5, 6 เป็นหน้าจอที่แสดงผลใน LINE app บนมือถือ (LIFF) รูปที่ 7, 8, 9 เป็นหน้าจอ desktop ผ่านเว็บไซต์ ความแตกต่างนี้สะท้อนให้เห็นว่าระบบรองรับการใช้งานทั้งบนมือถือและ desktop ในระบบเดียวกัน")

para()
para("เอกสารนี้จัดทำเพื่อรายงานผลการดำเนินงานต่อผู้บริหาร สำนักวิทยบริการ มหาวิทยาลัยนครพนม ปีงบประมาณ 2568",
     size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)

# ── Save ──────────────────────────────────────────────────────────────────────
out = "/mnt/c/projects/reserv/doc/report-management-2568.docx"
doc.save(out)
print(f"✓ Saved: {out}")
