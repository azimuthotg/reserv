"""สไตล์กลางสำหรับคู่มือชุด 2569 — ตัวหนังสือดำล้วน หัวตารางจัดกลางพื้นเทาอ่อน
ไม่ใช้ bullet และไม่มีกล่องคำแนะนำ/คำเตือน ตามรูปแบบที่ผู้ใช้กำหนด
"""
import os

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

FONT       = "TH Sarabun New"
BLACK      = RGBColor(0x00, 0x00, 0x00)
HEADER_BG  = "EDEDED"   # เทาอ่อนสำหรับหัวตาราง
BORDER_CLR = "BFBFBF"

BODY_PT    = 16
CONTENT_CM = 16.0       # ความกว้างพื้นที่พิมพ์ = 21 - (3 + 2)

TOC_INSTR  = ' TOC \\o "1-3" \\h \\z \\u '
TOC_TAB    = 9072       # จุดไข่ปลาชิดขวาที่ 16 ซม. (twips)


def el(tag, **attrs):
    """สร้าง element พร้อม attribute โดยเขียน prefix ด้วย `__` เช่น w__val='1'"""
    e = OxmlElement(tag)
    for k, v in attrs.items():
        e.set(qn(k.replace("__", ":")), v)
    return e


def ensure_toc_styles(doc):
    """เพิ่ม style toc 1–3 ให้ใช้ฟอนต์ไทยและจุดไข่ปลา

    Word สร้างย่อหน้าสารบัญด้วย style เหล่านี้ทับของเดิมตอนอัปเดต field
    ถ้าไม่กำหนดเอง ข้อความไทยจะกลับไปใช้ฟอนต์เริ่มต้นของ Word
    """
    styles = doc.styles.element
    existing = {s.get(qn("w:styleId")) for s in styles.findall(qn("w:style"))}

    for level, (indent, size) in enumerate(((0, 16), (567, 16), (1134, 15)), start=1):
        sid = f"TOC{level}"
        if sid in existing:
            continue

        st = el("w:style", w__type="paragraph", w__styleId=sid)
        st.append(el("w:name", w__val=f"toc {level}"))
        st.append(el("w:basedOn", w__val="Normal"))
        st.append(el("w:next", w__val="Normal"))
        st.append(el("w:uiPriority", w__val="39"))

        pPr = el("w:pPr")
        tabs = el("w:tabs")
        tabs.append(el("w:tab", w__val="right", w__leader="dot", w__pos=str(TOC_TAB)))
        pPr.append(tabs)
        pPr.append(el("w:spacing", w__after="60", w__line="276", w__lineRule="auto"))
        if indent:
            pPr.append(el("w:ind", w__left=str(indent)))
        st.append(pPr)

        rPr = el("w:rPr")
        rPr.append(el("w:rFonts", w__ascii=FONT, w__hAnsi=FONT, w__cs=FONT, w__eastAsia=FONT))
        rPr.append(el("w:color", w__val="000000"))
        rPr.append(el("w:sz", w__val=str(size * 2)))
        rPr.append(el("w:szCs", w__val=str(size * 2)))
        st.append(rPr)

        styles.append(st)


def wrap_toc_field(first_p, last_p):
    """ครอบย่อหน้าตั้งแต่ first_p ถึง last_p ด้วย TOC field ของ Word

    ย่อหน้าที่ครอบไว้กลายเป็น field result ชั่วคราว จึงยังอ่านได้ในโปรแกรมที่ไม่
    อัปเดต field ส่วน Word จะแทนที่ด้วยสารบัญจริงพร้อมเลขหน้าเมื่อเปิดไฟล์
    (ตั้ง w:dirty ไว้ ถ้าไม่ขึ้นให้กด Ctrl+A แล้ว F9)
    """
    begin = el("w:r")
    begin.append(el("w:fldChar", w__fldCharType="begin", w__dirty="true"))
    instr = el("w:r")
    it = el("w:instrText", xml__space="preserve")
    it.text = TOC_INSTR
    instr.append(it)
    sep = el("w:r")
    sep.append(el("w:fldChar", w__fldCharType="separate"))

    pPr = first_p._p.find(qn("w:pPr"))
    for node in (sep, instr, begin):     # แทรกย้อนกลับให้ได้ลำดับ begin → instr → sep
        if pPr is None:
            first_p._p.insert(0, node)
        else:
            pPr.addnext(node)

    end = el("w:r")
    end.append(el("w:fldChar", w__fldCharType="end"))
    last_p._p.append(end)


def _thai_run(element, size_pt, bold=False, italic=False):
    """ตั้งค่าฝั่ง Complex Script ให้ run หนึ่งตัว

    Word จัดข้อความไทยเป็น complex script จึงอ่านฟอนต์จาก w:cs และขนาดจาก w:szCs
    ไม่ใช่ w:ascii/w:sz ที่ python-docx ตั้งให้ ถ้าไม่ตั้งสองค่านี้ ข้อความไทยจะกลาย
    เป็นฟอนต์และขนาดเริ่มต้นของ Word ทั้งเอกสาร
    """
    rPr = element.get_or_add_rPr()

    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:cs"), FONT)
    rFonts.set(qn("w:eastAsia"), FONT)

    szCs = rPr.find(qn("w:szCs"))
    if szCs is None:
        szCs = OxmlElement("w:szCs")
        rPr.append(szCs)
    szCs.set(qn("w:val"), str(int(size_pt * 2)))

    if bold:
        b = OxmlElement("w:bCs")
        b.set(qn("w:val"), "1")
        rPr.append(b)
    if italic:
        i = OxmlElement("w:iCs")
        i.set(qn("w:val"), "1")
        rPr.append(i)


class Manual:
    def __init__(self, shots_dir):
        self.shots = shots_dir
        self.doc   = Document()
        self._setup()

    # ── ตั้งค่าเอกสาร ─────────────────────────────────────────────────────────
    def _setup(self):
        sec = self.doc.sections[0]
        sec.page_width  = Cm(21)
        sec.page_height = Cm(29.7)
        sec.top_margin    = Cm(2.5)
        sec.bottom_margin = Cm(2.0)
        sec.left_margin   = Cm(3.0)
        sec.right_margin  = Cm(2.0)

        normal = self.doc.styles["Normal"]
        normal.font.name = FONT
        normal.font.size = Pt(BODY_PT)
        normal.font.color.rgb = BLACK
        _thai_run(normal.element, BODY_PT)
        normal.paragraph_format.line_spacing = 1.15

        for level, size in ((1, 22), (2, 19), (3, 17)):
            s = self.doc.styles[f"Heading {level}"]
            s.font.name = FONT
            s.font.size = Pt(size)
            s.font.bold = True
            s.font.color.rgb = BLACK
            _thai_run(s.element, size, bold=True)

    # ── ชิ้นส่วนพื้นฐาน ───────────────────────────────────────────────────────
    def run(self, p, text, bold=False, size=BODY_PT, italic=False):
        r = p.add_run(text)
        r.font.name  = FONT
        r.font.size  = Pt(size)
        r.font.bold  = bold
        r.font.italic = italic
        r.font.color.rgb = BLACK
        _thai_run(r.element, size, bold, italic)
        return r

    def para(self, text="", bold=False, size=BODY_PT, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             before=0, after=6, indent=0.0, italic=False):
        p = self.doc.add_paragraph()
        p.alignment = align
        pf = p.paragraph_format
        pf.space_before = Pt(before)
        pf.space_after  = Pt(after)
        if indent:
            pf.left_indent = Cm(indent)
        if text:
            self.run(p, text, bold=bold, size=size, italic=italic)
        return p

    def heading(self, text, level=1, page_break_before=False):
        if page_break_before:
            self.doc.add_page_break()
        h = self.doc.add_heading(text, level=level)
        h.paragraph_format.space_before = Pt(16 if level == 1 else 12)
        h.paragraph_format.space_after  = Pt(8)
        size = {1: 22, 2: 19, 3: 17}.get(level, 17)
        for r in h.runs:
            r.font.name = FONT
            r.font.color.rgb = BLACK
            _thai_run(r.element, size, bold=True)
        return h

    def steps(self, items, indent=0.6):
        """ขั้นตอนแบบเลขลำดับ เริ่มที่ 1 เสมอ (ไม่ใช้ List Number ของ Word
        เพราะตัวนับใช้ร่วมกันทั้งเอกสาร)"""
        for i, text in enumerate(items, start=1):
            p = self.doc.add_paragraph()
            pf = p.paragraph_format
            pf.left_indent      = Cm(indent + 0.8)
            pf.first_line_indent = Cm(-0.8)
            pf.space_after      = Pt(4)
            self.run(p, f"{i}.", bold=True)
            self.run(p, "\t" + text)

    # ── ตาราง ────────────────────────────────────────────────────────────────
    def table(self, headers, rows, widths=None, center_cols=()):
        t = self.doc.add_table(rows=1 + len(rows), cols=len(headers))
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        self._borders(t)

        for i, h in enumerate(headers):
            cell = t.rows[0].cells[i]
            self._shade(cell, HEADER_BG)
            cp = cell.paragraphs[0]
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp.paragraph_format.space_before = Pt(3)
            cp.paragraph_format.space_after  = Pt(3)
            self.run(cp, h, bold=True, size=15)

        for ri, row_data in enumerate(rows):
            for ci, val in enumerate(row_data):
                cell = t.rows[ri + 1].cells[ci]
                cp = cell.paragraphs[0]
                cp.alignment = (WD_ALIGN_PARAGRAPH.CENTER if ci in center_cols
                                else WD_ALIGN_PARAGRAPH.LEFT)
                cp.paragraph_format.space_before = Pt(2)
                cp.paragraph_format.space_after  = Pt(2)
                self.run(cp, str(val), size=15)

        if widths:
            for ci, w in enumerate(widths):
                for row in t.rows:
                    row.cells[ci].width = Cm(w)

        self.para(after=4)
        return t

    def _shade(self, cell, fill):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill)
        tcPr.append(shd)

    def _borders(self, table):
        tbl = table._tbl
        tblPr = tbl.tblPr
        borders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            el = OxmlElement(f"w:{edge}")
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), "6")
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), BORDER_CLR)
            borders.append(el)
        tblPr.append(borders)

    # ── ภาพประกอบ 16:9 พร้อมคำบรรยาย ─────────────────────────────────────────
    def figure(self, filename, caption, width_cm=CONTENT_CM):
        path = filename if os.path.isabs(filename) else os.path.join(self.shots, filename)
        if not path.endswith(".png"):
            path += ".png"
        if not os.path.exists(path):
            print(f"  ! ไม่พบภาพ {path} — ข้าม")
            return False

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(2)
        p.add_run().add_picture(path, width=Cm(width_cm))

        self.figure_no = getattr(self, "figure_no", 0) + 1
        cp = self.doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(12)
        self.run(cp, f"ภาพที่ {self.figure_no}  {caption}", size=14)
        return True

    # ── ปก / สารบัญ ───────────────────────────────────────────────────────────
    def cover(self, title, subtitle, kind, org, version):
        self.para(before=90, after=0)
        self.para(title, bold=True, size=34, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
        self.para(subtitle, bold=True, size=22, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
        self.para(kind, size=18, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
        self.para(before=110, after=0)
        self.para(org, size=18, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
        self.para(version, size=16, align=WD_ALIGN_PARAGRAPH.CENTER)
        self.doc.add_page_break()

    def toc(self, entries):
        """สารบัญเป็น TOC field ของ Word — เลขหน้าเติมเองตอนเปิดไฟล์

        รายการที่ส่งเข้ามาใช้เป็นข้อความสำรองระหว่างที่ field ยังไม่ถูกอัปเดต
        จึงต้องเรียงให้ตรงกับหัวข้อจริงในเล่ม แต่ไม่ต้องใส่เลขหน้าเอง
        """
        ensure_toc_styles(self.doc)
        self.para("สารบัญ", bold=True, size=22, align=WD_ALIGN_PARAGRAPH.CENTER, after=14)

        paras = []
        for label, title in entries:
            p = self.doc.add_paragraph()
            p.paragraph_format.space_after = Pt(5)
            p.paragraph_format.left_indent = Cm(0.8 if label.count(".") else 0)
            self.run(p, f"{label}  {title}", bold=not label.count("."))
            paras.append(p)
        if paras:
            wrap_toc_field(paras[0], paras[-1])

        # section break แทน page break เพื่อให้ปก+สารบัญไม่มีเลขหน้า
        # และเนื้อหาเริ่มนับหน้าใหม่ที่ 1 (ตั้งค่าใน page_numbers)
        self.doc.add_section(WD_SECTION.NEW_PAGE)

    # ── หมายเลขหน้า ──────────────────────────────────────────────────────────
    def page_numbers(self):
        """ใส่เลขหน้าเฉพาะ section เนื้อหา และเริ่มนับที่ 1

        section แรก (ปก + สารบัญ) ไม่มี footer reference จึงไม่มีเลขหน้า
        """
        sec = self.doc.sections[-1]

        old = sec._sectPr.find(qn("w:pgNumType"))
        if old is not None:
            sec._sectPr.remove(old)
        anchor = None
        for tag in ("w:pgSz", "w:pgMar", "w:paperSrc", "w:pgBorders", "w:lnNumType"):
            found = sec._sectPr.find(qn(tag))
            if found is not None:
                anchor = found
        start = el("w:pgNumType", w__start="1")
        if anchor is None:
            sec._sectPr.insert(0, start)
        else:
            anchor.addnext(start)

        sec.footer.is_linked_to_previous = False
        p = sec.footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run()
        r.font.name = FONT
        r.font.size = Pt(14)
        r.font.color.rgb = BLACK
        _thai_run(r._r, 14)
        for tag, attrs, text in (
            ("w:fldChar", {"w:fldCharType": "begin"}, None),
            ("w:instrText", {"xml:space": "preserve"}, " PAGE "),
            ("w:fldChar", {"w:fldCharType": "end"}, None),
        ):
            node = OxmlElement(tag)
            for k, v in attrs.items():
                node.set(qn(k), v)
            if text:
                node.text = text
            r._r.append(node)

    def save(self, path):
        self.page_numbers()
        self.doc.save(path)
        print("บันทึกแล้ว:", path)
