"""Generate admin-manual.docx for the Staff Portal."""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TEAL = RGBColor(0x22, 0x97, 0x99)
GREEN = RGBColor(0x06, 0xC7, 0x55)
DARK = RGBColor(0x1a, 0x1f, 0x2e)
RED = RGBColor(0xdc, 0x35, 0x45)
GRAY = RGBColor(0x71, 0x80, 0x96)
LIGHT_BG = RGBColor(0xf8, 0xf9, 0xfa)
TEAL_BG = RGBColor(0xe6, 0xf4, 0xf4)
FONT = "TH Sarabun New"
SCREENSHOTS = "/mnt/c/projects/reserv/doc/screenshots"

doc = Document()

# ── Page setup ────────────────────────────────────────────────────────────────
sec = doc.sections[0]
sec.page_width  = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin    = Cm(2.5)
sec.bottom_margin = Cm(2.5)
sec.left_margin   = Cm(3.0)
sec.right_margin  = Cm(2.0)

# ── Default style ─────────────────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = FONT
style.font.size = Pt(14)
style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT)

def set_heading_style(level, size, color=DARK):
    name = f'Heading {level}'
    s = doc.styles[name]
    s.font.name = FONT
    s.font.size = Pt(size)
    s.font.bold = True
    s.font.color.rgb = color
    s.element.rPr.rFonts.set(qn('w:eastAsia'), FONT)

set_heading_style(1, 20, DARK)
set_heading_style(2, 17, TEAL)
set_heading_style(3, 15, DARK)

def run(para, text, bold=False, size=14, color=None, italic=False, underline=False):
    r = para.add_run(text)
    r.font.name = FONT
    r.font.bold = bold
    r.font.size = Pt(size)
    r.font.italic = italic
    r.font.underline = underline
    r.element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
    if color:
        r.font.color.rgb = color
    return r

def para(text="", bold=False, size=14, color=None, align=WD_ALIGN_PARAGRAPH.LEFT,
         space_before=0, space_after=6, italic=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        run(p, text, bold=bold, size=size, color=color, italic=italic)
    return p

def add_image(path, width_cm=14):
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(10)
        r = p.add_run()
        r.add_picture(path, width=Cm(width_cm))

def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    h.paragraph_format.space_after  = Pt(6)
    for r2 in h.runs:
        r2.font.name = FONT
        r2.element.rPr.rFonts.set(qn('w:eastAsia'), FONT)

def bullet(text, indent=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5 + indent * 0.5)
    p.paragraph_format.space_after = Pt(3)
    run(p, text, size=14)

def numbered(text, indent=0):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.left_indent = Cm(0.5 + indent * 0.5)
    p.paragraph_format.space_after = Pt(3)
    run(p, text, size=14)

def note_box(text):
    """Shaded note/caution box."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = t.cell(0, 0)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'FFF3CD')
    tcPr.append(shd)
    cp = cell.paragraphs[0]
    cp.paragraph_format.space_before = Pt(4)
    cp.paragraph_format.space_after = Pt(4)
    run(cp, "⚠ ", bold=True, size=13, color=RGBColor(0x85, 0x64, 0x04))
    run(cp, text, size=13, color=RGBColor(0x85, 0x64, 0x04))
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def info_box(text):
    """Teal info box."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = t.cell(0, 0)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'E6F4F4')
    tcPr.append(shd)
    cp = cell.paragraphs[0]
    cp.paragraph_format.space_before = Pt(4)
    cp.paragraph_format.space_after = Pt(4)
    run(cp, "ℹ ", bold=True, size=13, color=TEAL)
    run(cp, text, size=13, color=DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def simple_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Header row
    hrow = t.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1A1F2E')
        tcPr.append(shd)
        cp = cell.paragraphs[0]
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run(cp, h, bold=True, size=12, color=RGBColor(0xFF, 0xFF, 0xFF))
    # Data rows
    for ri, row_data in enumerate(rows):
        row = t.rows[ri+1]
        fill = 'F8F9FA' if ri % 2 == 0 else 'FFFFFF'
        for ci, cell_text in enumerate(row_data):
            cell = row.cells[ci]
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), fill)
            tcPr.append(shd)
            cp = cell.paragraphs[0]
            run(cp, str(cell_text), size=12)
    # Column widths
    if col_widths:
        for ci, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[ci].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return t

def page_break():
    doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════════════════════════════════
p = para()
p.paragraph_format.space_before = Pt(80)

p = para("Staff Portal", bold=True, size=36, color=TEAL, align=WD_ALIGN_PARAGRAPH.CENTER)
p = para("Library@NPU", bold=True, size=28, color=DARK, align=WD_ALIGN_PARAGRAPH.CENTER)
p = para("คู่มือผู้ดูแลระบบ", bold=True, size=22, color=DARK, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=20)
p = para("ระบบจองพื้นที่บริการ Smart Creative Learning Space", size=16, align=WD_ALIGN_PARAGRAPH.CENTER)

p = para()
p.paragraph_format.space_before = Pt(60)

p = para("สำนักวิทยบริการ มหาวิทยาลัยนครพนม", size=15, align=WD_ALIGN_PARAGRAPH.CENTER)
p = para("เวอร์ชัน 1.0  |  เมษายน 2569", size=13, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# บทที่ 1 บทนำ
# ═══════════════════════════════════════════════════════════════════════════════
heading("บทที่ 1  บทนำ", 1)

heading("1.1  ภาพรวม Staff Portal", 2)
para("Staff Portal คือส่วนจัดการระบบจองพื้นที่บริการ Smart Creative Learning Space สำหรับเจ้าหน้าที่และผู้ดูแลระบบของสำนักวิทยบริการ ใช้สำหรับ:")
bullet("ติดตามรายการจองและยกเลิกการจองในกรณีที่จำเป็น")
bullet("กำหนดวันหยุดเพื่อปิดรับการจองอัตโนมัติ")
bullet("ตรวจสอบประวัติและจัดการสิทธิ์ผู้ใช้ LINE")
bullet("จัดการข้อมูลห้องและบัญชีเจ้าหน้าที่ (เฉพาะผู้ดูแลระบบ)")

heading("1.2  บทบาทและสิทธิ์ผู้ใช้", 2)
para("ระบบมี 2 บทบาทที่มีสิทธิ์การเข้าถึงต่างกัน:")
simple_table(
    ["เมนู", "เจ้าหน้าที่", "ผู้ดูแลระบบ"],
    [
        ["Dashboard", "✅", "✅"],
        ["รายการจอง", "✅", "✅"],
        ["วันหยุด", "✅", "✅"],
        ["ผู้ใช้ LINE", "✅", "✅"],
        ["จัดการห้อง", "❌", "✅"],
        ["จัดการเจ้าหน้าที่", "❌", "✅"],
    ],
    col_widths=[9, 3, 3]
)
info_box("ผู้ดูแลระบบ (Admin) จะเห็นเมนูเพิ่มเติมในหัวข้อ 'ผู้ดูแลระบบ' ที่ Sidebar ส่วนเจ้าหน้าที่ทั่วไปจะไม่เห็นเมนูดังกล่าว")

heading("1.3  วิธีเข้าถึงระบบ", 2)
para("เข้าระบบผ่านเบราว์เซอร์ที่ URL:")
p = para("https://lib.npu.ac.th/reserv/manage/", bold=True, size=14, color=TEAL)
para("รองรับเบราว์เซอร์: Google Chrome, Microsoft Edge, Firefox (รุ่นล่าสุด)")

heading("1.4  การรองรับอุปกรณ์", 2)
bullet("คอมพิวเตอร์ — ใช้งานได้เต็มประสิทธิภาพ ขนาดหน้าจอแนะนำ 1280 px ขึ้นไป")
bullet("แท็บเล็ต / มือถือ — รองรับ Responsive Design แต่แนะนำให้ใช้คอมพิวเตอร์เพื่อประสิทธิภาพสูงสุด")

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# บทที่ 2 การเข้าและออกจากระบบ
# ═══════════════════════════════════════════════════════════════════════════════
heading("บทที่ 2  การเข้าและออกจากระบบ", 1)

heading("2.1  เข้าสู่ระบบ", 2)
add_image(f"{SCREENSHOTS}/admin_01_login.png", width_cm=10)

numbered("เปิดเบราว์เซอร์แล้วไปที่ https://lib.npu.ac.th/reserv/manage/")
numbered("กรอก ชื่อผู้ใช้ (username) ในช่องแรก")
numbered("กรอก รหัสผ่าน (password) ในช่องที่สอง")
numbered('คลิกปุ่ม "เข้าสู่ระบบ"')

para("หากล็อกอินสำเร็จ ระบบจะพาไปยังหน้า Dashboard โดยอัตโนมัติ")
note_box("หากกรอกชื่อผู้ใช้หรือรหัสผ่านผิด จะแสดงข้อความ 'ชื่อผู้ใช้หรือรหัสผ่านไม่ถูกต้อง หรือไม่มีสิทธิ์เข้าใช้' บัญชีที่ไม่ใช่ Staff หรือ Admin จะเข้าระบบนี้ไม่ได้")

heading("2.2  หน้าตาหลัก — Sidebar และ Topbar", 2)
para("เมื่อเข้าสู่ระบบแล้วจะพบส่วนประกอบหลัก 2 ส่วน:")
bullet("Sidebar (แถบซ้าย) — เมนูนำทางหลัก พื้นหลังสีเข้ม แสดงชื่อ 'Staff Portal — Library@NPU' และรายการเมนูทั้งหมด เมนูที่กำลังใช้งานจะไฮไลต์ด้วยสีเขียว")
bullet("Topbar (แถบบน) — แสดงชื่อหน้าปัจจุบัน และบทบาทของผู้ใช้ที่ล็อกอินอยู่ (เจ้าหน้าที่ / ผู้ดูแลระบบ)")

heading("2.3  ออกจากระบบ", 2)
para("คลิก 'ออกจากระบบ' ที่ด้านล่างของ Sidebar ระบบจะพาไปยังหน้าล็อกอินทันที")

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# บทที่ 3 Dashboard
# ═══════════════════════════════════════════════════════════════════════════════
heading("บทที่ 3  Dashboard", 1)
para("เข้าถึงได้โดยคลิก 'Dashboard' ใน Sidebar")
add_image(f"{SCREENSHOTS}/admin_02_dashboard.png", width_cm=15)
para("Dashboard คือหน้าแรกหลังเข้าสู่ระบบ แสดงภาพรวมสถานการณ์การจองในปัจจุบัน")

heading("3.1  การ์ดสถิติ", 2)
para("ด้านบนของหน้ามีการ์ดสถิติ 4 ใบ:")
simple_table(
    ["การ์ด", "ความหมาย"],
    [
        ["การจองวันนี้", "จำนวนการจองที่ยืนยันแล้วในวันนี้"],
        ["รอดำเนินการ", "จำนวนการจองที่ยืนยันแล้วตั้งแต่วันนี้เป็นต้นไป"],
        ["จองเดือนนี้", "จำนวนการจองทั้งหมดในเดือนปัจจุบัน"],
        ["ผู้ใช้ LINE", "จำนวนผู้ใช้ที่ลงทะเบียนและเปิดใช้งานอยู่"],
    ],
    col_widths=[5, 10]
)

heading("3.2  ตารางการจองล่าสุด", 2)
para("แสดงรายการจอง 10 รายการล่าสุดที่มีสถานะ 'ยืนยันแล้ว' ประกอบด้วย ห้อง, วันที่, เวลา, ชื่อกลุ่ม และผู้จอง")
para("คลิก 'ดูทั้งหมด' เพื่อไปยังหน้ารายการจองครบถ้วน")

heading("3.3  วันหยุดที่กำลังจะมา", 2)
para("แสดงวันหยุดที่เปิดใช้งานในช่วง 30 วันข้างหน้า คลิก 'จัดการ' เพื่อไปยังหน้าจัดการวันหยุด")

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# บทที่ 4 จัดการรายการจอง
# ═══════════════════════════════════════════════════════════════════════════════
heading("บทที่ 4  จัดการรายการจอง", 1)
para("เข้าถึงได้โดยคลิก 'รายการจอง' ใน Sidebar  |  สิทธิ์: เจ้าหน้าที่และผู้ดูแลระบบ")
add_image(f"{SCREENSHOTS}/admin_03_bookings.png", width_cm=15)

heading("4.1  ดูและค้นหารายการจอง", 2)
para("ด้านบนของหน้ามีแถบกรองข้อมูล 4 ช่อง:")
simple_table(
    ["ช่อง", "ใช้สำหรับ"],
    [
        ["ห้อง", "กรองแสดงเฉพาะห้องที่เลือก"],
        ["วันที่", "กรองตามวันที่จอง"],
        ["สถานะ", "เลือกระหว่าง ทั้งหมด / ยืนยันแล้ว / ยกเลิกแล้ว"],
        ["ค้นหา", "พิมพ์ชื่อกลุ่ม, ชื่อผู้จอง หรือชื่อคณะ"],
    ],
    col_widths=[4, 11]
)
numbered("ตั้งค่า filter ที่ต้องการ")
numbered("คลิกปุ่ม 'ค้นหา'")
numbered("ระบบแสดงผลและบอกจำนวนรายการที่พบ")

para("คอลัมน์ในตาราง:")
simple_table(
    ["คอลัมน์", "ความหมาย"],
    [
        ["#", "หมายเลขการจอง"],
        ["ห้อง", "ชื่อห้องที่จอง"],
        ["วันที่", "วันที่จอง"],
        ["เวลา", "เวลาเริ่ม–สิ้นสุด"],
        ["กลุ่ม / กิจกรรม", "ชื่อกลุ่ม (ถ้ายกเลิก จะแสดงเหตุผลสีแดงด้านล่าง)"],
        ["ผู้จอง", "ชื่อผู้จอง (คลิกได้เพื่อดูประวัติ)"],
        ["คณะ", "คณะหรือหน่วยงานของผู้จอง"],
        ["สถานะ", "ยืนยันแล้ว (สีเขียว) / ยกเลิก (สีแดง)"],
    ],
    col_widths=[5, 10]
)
info_box("หากมีข้อมูลเกิน 20 รายการ ระบบจะแบ่งหน้า (Pagination) กดลูกศรซ้าย–ขวาที่ด้านล่างตารางเพื่อเลื่อนหน้า")

heading("4.2  ยกเลิกการจอง", 2)
add_image(f"{SCREENSHOTS}/admin_04_cancel_modal.png", width_cm=12)
numbered("หาการจองที่ต้องการยกเลิกในตาราง")
numbered("คลิกปุ่ม 'ยกเลิก' (สีแดง) ที่ด้านขวาของแถว")
numbered("กล่องยืนยัน (Modal) จะเปิดขึ้น แสดงชื่อกลุ่มที่จะยกเลิก")
numbered("กรอก เหตุผลการยกเลิก (ไม่บังคับ แต่แนะนำ เช่น 'ห้องปิดปรับปรุง')")
numbered("คลิกปุ่ม 'ยืนยันยกเลิก' (สีแดง)")

info_box("ผลที่เกิดขึ้น: สถานะการจองเปลี่ยนเป็น 'ยกเลิก' และระบบส่ง LINE Notification แจ้งผู้จองอัตโนมัติพร้อมเหตุผล (ถ้ามี)")

heading("4.3  ดู Logs การจอง", 2)
add_image(f"{SCREENSHOTS}/admin_05_booking_logs.png", width_cm=15)
numbered("คลิกไอคอนนาฬิกา 🕐 ที่ด้านขวาของแถวการจอง")
numbered("หน้า Logs จะแสดงข้อมูลสรุปการจอง และ Timeline ประวัติการดำเนินการ")

para("รายการ Action ใน Timeline:")
simple_table(
    ["Action", "ไอคอน", "สี", "ความหมาย"],
    [
        ["สร้างการจอง", "➕", "เขียว", "ผู้ใช้จองสำเร็จ"],
        ["ยกเลิกการจอง", "✖", "แดง", "ยกเลิกโดยผู้ใช้หรือเจ้าหน้าที่"],
        ["ปิดอุปกรณ์อัตโนมัติ", "⚡", "เหลือง", "ระบบปิดอุปกรณ์ IoT หลังหมดเวลา"],
        ["เข้าถึงข้อมูล", "ℹ", "น้ำเงิน", "เจ้าหน้าที่ตรวจสอบรายการ"],
    ],
    col_widths=[4, 2, 2, 7]
)

heading("4.4  ดูโปรไฟล์ผู้จอง", 2)
para("ในหน้ารายการจอง คลิกที่ชื่อผู้จองเพื่อดูข้อมูลส่วนตัวและประวัติการจองทั้งหมด (ดูรายละเอียดในบทที่ 6.2)")

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# บทที่ 5 จัดการวันหยุด
# ═══════════════════════════════════════════════════════════════════════════════
heading("บทที่ 5  จัดการวันหยุด", 1)
para("เข้าถึงได้โดยคลิก 'วันหยุด' ใน Sidebar  |  สิทธิ์: เจ้าหน้าที่และผู้ดูแลระบบ")
add_image(f"{SCREENSHOTS}/admin_06_holidays.png", width_cm=15)
para("เมื่อกำหนดวันหยุด ระบบจะปิดรับการจองอัตโนมัติในวันดังกล่าว ผู้ใช้จะไม่สามารถเลือกวันนั้นในปฏิทินได้")

heading("5.1  ดูรายการวันหยุด", 2)
bullet("ระบบแสดงวันหยุดแยกตามปี")
bullet("คลิกปุ่มปี (2026 / 2027 / 2028) เพื่อสลับดูข้อมูลของแต่ละปี")
bullet("แต่ละรายการแสดง: วันที่, วันในสัปดาห์, คำอธิบาย, สถานะ (เปิดใช้ / ปิดใช้)")

heading("5.2  เพิ่มวันหยุด", 2)
add_image(f"{SCREENSHOTS}/admin_07_holiday_form.png", width_cm=12)
numbered("คลิกปุ่ม '+ เพิ่มวันหยุด' มุมบนขวา")
numbered("กรอกข้อมูล: วันที่, คำอธิบาย (ชื่อวันหยุด), ติ๊ก 'เปิดใช้งาน' ถ้าต้องการปิดรับจองทันที")
numbered("คลิกปุ่ม 'บันทึก'")

heading("5.3  แก้ไขวันหยุด", 2)
numbered("คลิกปุ่ม 'แก้ไข' ที่แถววันหยุดที่ต้องการ")
numbered("แก้ไขข้อมูลในฟอร์ม")
numbered("คลิกปุ่ม 'บันทึก'")

heading("5.4  เปิด / ปิดใช้งานวันหยุด", 2)
para("ใช้สำหรับกรณีที่ต้องการ 'พักใช้' วันหยุดชั่วคราวโดยไม่ลบออก:")
bullet("คลิก 'ปิดใช้' (สีเหลือง) — วันหยุดยังอยู่ในระบบแต่รับจองได้ตามปกติ")
bullet("คลิก 'เปิดใช้' (สีเขียว) — วันหยุดกลับมาใช้งาน ระบบปิดรับจองวันนั้นทันที")

heading("5.5  ลบวันหยุด", 2)
numbered("คลิกปุ่ม 'ลบ' (สีแดง) ที่แถววันหยุดที่ต้องการ")
numbered("กล่องยืนยันจะเปิดขึ้น แสดงชื่อวันหยุดที่จะลบ")
numbered("คลิก 'ลบ' เพื่อยืนยัน")
note_box("การลบวันหยุดจะนำวันนั้นออกจากระบบถาวร หากต้องการเก็บข้อมูลไว้ แนะนำให้ใช้การ 'ปิดใช้' แทน")

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# บทที่ 6 จัดการผู้ใช้ LINE
# ═══════════════════════════════════════════════════════════════════════════════
heading("บทที่ 6  จัดการผู้ใช้ LINE", 1)
para("เข้าถึงได้โดยคลิก 'ผู้ใช้ LINE' ใน Sidebar  |  สิทธิ์: เจ้าหน้าที่และผู้ดูแลระบบ")
add_image(f"{SCREENSHOTS}/admin_08_line_users.png", width_cm=15)
para("หน้านี้แสดงรายชื่อผู้ใช้ทุกคนที่ลงทะเบียนผูกบัญชี LINE กับระบบแล้ว")

heading("6.1  ดูรายชื่อผู้ใช้", 2)
simple_table(
    ["คอลัมน์", "ความหมาย"],
    [
        ["ชื่อ-นามสกุล", "ชื่อจริงจาก NPU API (บรรทัดล่างคือชื่อ LINE)"],
        ["รหัส LDAP", "รหัสประจำตัวนักศึกษา / บุคลากร"],
        ["ประเภท", "นักศึกษา หรือ บุคลากรภายในมหาวิทยาลัย"],
        ["คณะ / หน่วยงาน", "สังกัด"],
        ["สาขา / ฝ่าย", "สาขาวิชา (นักศึกษา) หรือฝ่าย (บุคลากร)"],
        ["วันที่ลงทะเบียน", "วันแรกที่ผูกบัญชี"],
        ["สถานะ", "ใช้งาน (เขียว) / ปิดใช้ (แดง)"],
    ],
    col_widths=[5, 10]
)

para("การค้นหา:")
numbered("กรอกคำค้นในช่อง 'ค้นหา' (ค้นตามชื่อ, รหัส LDAP หรือคณะ)")
numbered("เลือก 'ประเภท' เพื่อกรองเฉพาะนักศึกษาหรือบุคลากร")
numbered("คลิก 'ค้นหา'")

heading("6.2  ดูโปรไฟล์และประวัติการจองรายบุคคล", 2)
add_image(f"{SCREENSHOTS}/admin_09_user_detail.png", width_cm=15)
numbered("คลิกไอคอนรายชื่อ 👤 หรือคลิกชื่อผู้ใช้ในตาราง")
numbered("หน้าโปรไฟล์จะแสดงข้อมูลดังนี้:")
bullet("ข้อมูลส่วนตัว — ชื่อ-นามสกุล, ชื่อ LINE, ประเภทผู้ใช้, สถานะ, รหัส LDAP, คณะ/หน่วยงาน, สาขา/ฝ่าย, วันที่ลงทะเบียน", indent=1)
bullet("สถิติการจอง — จองทั้งหมด, ยืนยันแล้ว, ยกเลิกแล้ว, อัตราการยกเลิก (%)", indent=1)
bullet("ประวัติการจองทั้งหมด — ตารางแสดงการจองทุกครั้ง พร้อมเหตุผลยกเลิก (ถ้ามี) และปุ่ม Logs", indent=1)

heading("6.3  เปิด / ปิดใช้งานผู้ใช้", 2)
para("เมื่อปิดใช้งานผู้ใช้ ผู้ใช้รายนั้นจะไม่สามารถเข้าใช้ระบบจองได้จนกว่าจะเปิดใช้งานอีกครั้ง")
bullet("จากหน้ารายชื่อ — คลิกปุ่ม 'ปิดใช้' (สีเหลือง) หรือ 'เปิดใช้' (สีเขียว) ที่แถวของผู้ใช้นั้น")
bullet("จากหน้าโปรไฟล์ — คลิกปุ่ม 'ปิดใช้งาน' หรือ 'เปิดใช้งาน' มุมบนขวา")
info_box("แถวของผู้ใช้ที่ถูกปิดใช้งานจะแสดงสีจาง (โปร่งแสง 55%) เพื่อแยกให้เห็นชัดเจน")

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# บทที่ 7 จัดการห้อง
# ═══════════════════════════════════════════════════════════════════════════════
heading("บทที่ 7  จัดการห้อง", 1)
para("เข้าถึงได้โดยคลิก 'จัดการห้อง' ใน Sidebar (หัวข้อ 'ผู้ดูแลระบบ')  |  สิทธิ์: เฉพาะผู้ดูแลระบบ (Admin)")
add_image(f"{SCREENSHOTS}/admin_10_rooms.png", width_cm=15)

heading("7.1  ดูรายการห้องทั้งหมด", 2)
simple_table(
    ["คอลัมน์", "ความหมาย"],
    [
        ["ชื่อห้อง", "ชื่อเต็มที่แสดงในระบบ"],
        ["Booking Key", "รหัสย่อที่ใช้ใน URL เช่น mini, netflix"],
        ["ที่ตั้ง", "สถานที่ตั้งของห้อง"],
        ["ความจุ", "จำนวนคนสูงสุดที่รับได้"],
        ["เวลาเปิด–ปิด", "ช่วงเวลาที่รับการจอง"],
        ["สถานะ", "เปิด (เขียว) / ปิด (แดง)"],
    ],
    col_widths=[5, 10]
)

heading("7.2  เพิ่มห้องใหม่", 2)
numbered("คลิกปุ่ม '+ เพิ่มห้อง' มุมบนขวา")
numbered("กรอกข้อมูลห้อง: ชื่อห้อง, Booking Key, คำอธิบาย (ไม่บังคับ), ที่ตั้ง, ความจุ, เวลาเปิด–ปิด")
numbered("คลิก 'บันทึก'")
note_box("Booking Key ต้องเป็นภาษาอังกฤษตัวพิมพ์เล็ก ไม่มีช่องว่าง และไม่ซ้ำกับห้องอื่น เพราะจะใช้เป็น URL เช่น /booking/?room=mini")

heading("7.3  แก้ไขข้อมูลห้อง", 2)
numbered("คลิกปุ่ม 'แก้ไข' ที่แถวห้องที่ต้องการ")
numbered("แก้ไขข้อมูลในฟอร์ม")
numbered("คลิก 'บันทึก'")

heading("7.4  เปิด / ปิดห้อง", 2)
para("ใช้สำหรับปิดรับจองชั่วคราว เช่น ห้องปรับปรุง:")
bullet("คลิก 'ปิดห้อง' (สีเหลือง) — ห้องจะไม่ปรากฏในหน้าจองของผู้ใช้")
bullet("คลิก 'เปิดห้อง' (สีเขียว) — ห้องกลับมารับการจองได้ตามปกติ")

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# บทที่ 8 จัดการเจ้าหน้าที่
# ═══════════════════════════════════════════════════════════════════════════════
heading("บทที่ 8  จัดการเจ้าหน้าที่", 1)
para("เข้าถึงได้โดยคลิก 'เจ้าหน้าที่' ใน Sidebar (หัวข้อ 'ผู้ดูแลระบบ')  |  สิทธิ์: เฉพาะผู้ดูแลระบบ (Admin)")
add_image(f"{SCREENSHOTS}/admin_11_staff.png", width_cm=15)

heading("8.1  ดูรายชื่อเจ้าหน้าที่", 2)
para("ตารางแสดงบัญชีเจ้าหน้าที่และผู้ดูแลระบบทั้งหมด พร้อมข้อมูล:")
bullet("ชื่อผู้ใช้, ชื่อ-นามสกุล, อีเมล")
bullet("บทบาท — ผู้ดูแลระบบ (badge สีแดง) / เจ้าหน้าที่ (badge สีเทา)")
bullet("เวลาเข้าใช้ล่าสุด")
bullet("สถานะ — ใช้งาน (เขียว) / ปิดใช้ (แดง)")

heading("8.2  เพิ่มบัญชีเจ้าหน้าที่", 2)
numbered("คลิกปุ่ม '+ เพิ่มเจ้าหน้าที่'")
numbered("กรอกข้อมูล: ชื่อผู้ใช้ (ไม่ซ้ำ), ชื่อ-นามสกุล (ไม่บังคับ), อีเมล (ไม่บังคับ), รหัสผ่าน, ยืนยันรหัสผ่าน")
numbered("ติ๊ก 'บทบาทผู้ดูแลระบบ' ถ้าต้องการให้มีสิทธิ์ Admin เต็ม")
numbered("คลิก 'บันทึก'")
info_box("บัญชีใหม่จะเป็น Staff (เจ้าหน้าที่) โดยอัตโนมัติ เว้นแต่ติ๊ก 'บทบาทผู้ดูแลระบบ'")

heading("8.3  แก้ไขข้อมูล / เปลี่ยนรหัสผ่าน / ปรับสิทธิ์", 2)
numbered("คลิกปุ่ม 'แก้ไข' ที่แถวเจ้าหน้าที่ที่ต้องการ")
numbered("แก้ไขข้อมูล: ชื่อ, นามสกุล, อีเมล, สิทธิ์ (Admin ↔ Staff)")
numbered("หากต้องการเปลี่ยนรหัสผ่าน — กรอกรหัสใหม่ในช่อง 'รหัสผ่านใหม่' (ถ้าไม่กรอกจะไม่เปลี่ยน)")
numbered("คลิก 'บันทึก'")

heading("8.4  เปิด / ปิดใช้งานบัญชี", 2)
bullet("คลิก 'ปิดใช้' — บัญชีนั้นจะเข้าสู่ระบบไม่ได้ชั่วคราว")
bullet("คลิก 'เปิดใช้' — บัญชีกลับมาใช้งานได้")
note_box("ผู้ดูแลระบบไม่สามารถปิดบัญชีของตัวเองได้ ปุ่ม 'ปิดใช้' จะไม่ปรากฏในแถวของบัญชีที่กำลังล็อกอินอยู่")

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# บทที่ 9 FAQ
# ═══════════════════════════════════════════════════════════════════════════════
heading("บทที่ 9  คำถามที่พบบ่อย", 1)

heading("ลืมรหัสผ่าน ต้องทำอย่างไร?", 2)
para("ระบบไม่มีการรีเซ็ตรหัสผ่านอัตโนมัติ ต้องติดต่อผู้ดูแลระบบ (Admin) เพื่อเปลี่ยนรหัสผ่านในหน้า 'จัดการเจ้าหน้าที่ → แก้ไข'")

heading("ผู้ใช้ LINE ถูกปิดใช้งาน จะเปิดอีกครั้งได้อย่างไร?", 2)
numbered("ไปที่ ผู้ใช้ LINE")
numbered("ค้นหาชื่อหรือรหัส LDAP")
numbered("คลิกปุ่ม 'เปิดใช้' ที่แถวของผู้ใช้นั้น")

heading("ต้องการปิดห้องชั่วคราว ต้องทำอย่างไร?", 2)
para("กรณีปิดห้องทั้งวัน (เช่น ซ่อมแซม):")
bullet("ไปที่ จัดการห้อง (ต้องเป็น Admin) → คลิก 'ปิดห้อง' → เมื่อพร้อมคลิก 'เปิดห้อง'")
para("กรณีปิดเฉพาะวัน (เช่น งานพิเศษ):")
bullet("ไปที่ วันหยุด → เพิ่มวันที่นั้นเป็นวันหยุด")

heading("เพิ่มวันหยุดแล้วแต่ผู้ใช้ยังเลือกวันนั้นได้อยู่?", 2)
para("ตรวจสอบว่าสถานะของวันหยุดนั้นเป็น 'เปิดใช้' หรือไม่ ถ้าเป็น 'ปิดใช้' ระบบจะยังรับการจองวันนั้น คลิก 'เปิดใช้' เพื่อเปิดใช้งานวันหยุดนั้น")

heading("ต้องการดูว่าห้องไหนจองเยอะที่สุดในเดือนนี้?", 2)
para("ใช้หน้ารายการจอง กรองด้วย สถานะ = ยืนยันแล้ว จากนั้นเลือกห้องทีละห้องแล้วดูจำนวนรายการที่พบ")

heading("มีการจองซ้อนกัน (Conflict) จะรู้ได้อย่างไร?", 2)
para("ระบบตรวจสอบและป้องกันการจองซ้อนกันอัตโนมัติ ผู้ใช้จะไม่สามารถจองช่วงเวลาที่ถูกจองแล้วได้ เจ้าหน้าที่ไม่ต้องตรวจสอบด้วยตนเอง")

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# ภาคผนวก
# ═══════════════════════════════════════════════════════════════════════════════
heading("ภาคผนวก ก — ข้อมูลห้องบริการทั้ง 5 ห้อง", 1)
simple_table(
    ["ชื่อห้อง", "Booking Key", "ที่ตั้ง", "ความจุ", "เวลาเปิด–ปิด"],
    [
        ["Mini Theater", "mini", "ชั้น 1 อาคารบรรณสาร", "30 คน", "08:00–20:00"],
        ["Netflix Room", "netflix", "ชั้น 2 อาคารบรรณสาร", "10 คน", "08:00–20:00"],
        ["Canva Studio", "canva", "ชั้น 2 อาคารบรรณสาร", "8 คน", "08:00–18:00"],
        ["ChatGPT Room", "chat-gpt", "ชั้น 2 อาคารบรรณสาร", "6 คน", "08:00–18:00"],
        ["Meeting Room F1", "meeting_f1", "ชั้น 1 อาคารบรรณสาร", "20 คน", "08:00–20:00"],
    ],
    col_widths=[4, 3, 5, 2.5, 3]
)

heading("ภาคผนวก ข — การแจ้งเตือน LINE ที่ระบบส่งอัตโนมัติ", 1)
simple_table(
    ["เหตุการณ์", "ผู้รับแจ้งเตือน", "เนื้อหา"],
    [
        ["จองสำเร็จ", "ผู้จอง", "ชื่อห้อง, วันที่, เวลา, ชื่อกลุ่ม"],
        ["ยกเลิกโดยผู้ใช้", "ผู้จอง", "แจ้งการยกเลิกพร้อมเหตุผล"],
        ["ยกเลิกโดยเจ้าหน้าที่", "ผู้จอง", "แจ้งการยกเลิกพร้อมเหตุผล (ถ้ากรอก)"],
    ],
    col_widths=[5, 3.5, 7]
)
info_box("เจ้าหน้าที่ไม่ต้องแจ้งผู้จองด้วยตนเอง ระบบจัดการส่งแจ้งเตือนให้อัตโนมัติทุกครั้ง")

# ── Footer note ───────────────────────────────────────────────────────────────
para()
p = para("คู่มือนี้จัดทำโดยสำนักวิทยบริการ มหาวิทยาลัยนครพนม",
         size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)
p = para("Smart Creative Learning Space  |  เวอร์ชัน 1.0  |  เมษายน 2569",
         size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)

# ── Save ──────────────────────────────────────────────────────────────────────
out_path = "/mnt/c/projects/reserv/doc/admin-manual.docx"
doc.save(out_path)
print(f"✓ Saved: {out_path}")
