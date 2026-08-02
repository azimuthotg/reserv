"""แพตช์คู่มือ .docx ที่สร้างไว้แล้ว ให้สารบัญมีเลขหน้า และเริ่มนับหน้า 1 ที่บทที่ 1

ใช้กับไฟล์ที่แก้ด้วย Word ไปแล้ว จึงไม่รันสคริปต์สร้างใหม่ (งานที่แก้มือจะหาย)
สิ่งที่สคริปต์นี้ทำกับไฟล์เดิม:

1. ครอบรายการสารบัญเดิมด้วย TOC field ของ Word (`TOC \\o "1-3" \\h \\z \\u`)
   รายการเดิมกลายเป็น field result ชั่วคราว จึงยังอ่านได้ถ้าเปิดด้วยโปรแกรมที่
   ไม่อัปเดต field ส่วน Word จะแทนที่ด้วยสารบัญจริงพร้อมเลขหน้าเมื่อเปิดไฟล์
   (ตั้ง w:dirty ไว้แล้ว ถ้าไม่ขึ้นให้กด Ctrl+A แล้ว F9)
2. เพิ่ม style toc 1–3 ให้ใช้ฟอนต์ไทยและจุดไข่ปลาชิดขวา เพราะ Word จะสร้าง
   ย่อหน้าสารบัญด้วย style เหล่านี้ทับของเดิม
3. เปลี่ยน page break หลังสารบัญเป็น section break แล้วให้ section แรก
   (ปก + สารบัญ) ไม่มี footer ส่วน section เนื้อหาเริ่มนับหน้าใหม่ที่ 1

    python3 doc/fix_manual_toc.py doc/user-manual-reserv-2569.docx
"""
import copy
import os
import shutil
import sys

from docx import Document
from docx.oxml.ns import qn

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from manual_style import ensure_toc_styles, el as _el, wrap_toc_field  # noqa: E402


# ── 1. ครอบสารบัญเดิมด้วย TOC field ──────────────────────────────────────────
def add_toc_field(doc):
    """หารายการสารบัญเดิมแล้วครอบด้วย TOC field คืนค่า index ของหัวข้อ 'สารบัญ'"""
    paras = doc.paragraphs
    start = next((i for i, p in enumerate(paras) if p.text.strip() == "สารบัญ"), None)
    if start is None:
        raise SystemExit("ไม่พบหัวข้อ 'สารบัญ' ในเอกสาร")

    if any("TOC \\o" in p._p.xml for p in paras[start:start + 40]):
        print("  - มี TOC field อยู่แล้ว ข้ามขั้นตอนนี้")
        return start

    entries = []
    for p in paras[start + 1:]:
        if p.style.name.startswith("Heading") or "w:br" in p._p.xml:
            break
        if p.text.strip():
            entries.append(p)
    if not entries:
        raise SystemExit("ไม่พบรายการสารบัญหลังหัวข้อ 'สารบัญ'")

    wrap_toc_field(entries[0], entries[-1])
    print(f"  - ครอบ TOC field รอบรายการสารบัญ {len(entries)} บรรทัด")
    return start


# ── 3. section break + เริ่มนับหน้าใหม่ ───────────────────────────────────────
def _insert_ordered(sectPr, node, after_tags):
    """แทรก node ตามลำดับ schema ของ CT_SectPr"""
    prev = None
    for tag in after_tags:
        found = sectPr.find(qn(tag))
        if found is not None:
            prev = found
    if prev is None:
        sectPr.insert(0, node)
    else:
        prev.addnext(node)


def split_and_number(doc, toc_index):
    body_sectPr = doc.sections[-1]._sectPr

    if len(doc.sections) > 1:
        print("  - เอกสารมี section break อยู่แล้ว ข้ามการแบ่ง section")
    else:
        paras = doc.paragraphs
        brk = next((p for p in paras[toc_index:] if "w:br" in p._p.xml), None)
        if brk is None:
            raise SystemExit("ไม่พบ page break หลังสารบัญ")

        for r in list(brk._p.findall(qn("w:r"))):
            brk._p.remove(r)

        sect = copy.deepcopy(body_sectPr)
        for fref in sect.findall(qn("w:footerReference")):   # ปก+สารบัญ ไม่มีเลขหน้า
            sect.remove(fref)
        for href in sect.findall(qn("w:headerReference")):
            sect.remove(href)
        for t in sect.findall(qn("w:type")):
            sect.remove(t)
        _insert_ordered(sect, _el("w:type", w__val="nextPage"),
                        ("w:footerReference", "w:headerReference",
                         "w:footnotePr", "w:endnotePr"))

        pPr = brk._p.get_or_add_pPr()
        pPr.append(sect)
        print("  - เปลี่ยน page break หลังสารบัญเป็น section break (ปก+สารบัญ ไม่มีเลขหน้า)")

    for old in body_sectPr.findall(qn("w:pgNumType")):
        body_sectPr.remove(old)
    _insert_ordered(body_sectPr, _el("w:pgNumType", w__start="1"),
                    ("w:pgSz", "w:pgMar", "w:paperSrc", "w:pgBorders", "w:lnNumType"))
    print("  - ตั้งให้ section เนื้อหาเริ่มนับหน้าที่ 1")


def main(path):
    if not os.path.exists(path):
        raise SystemExit(f"ไม่พบไฟล์ {path}")

    backup = path.replace(".docx", ".before-toc.docx")
    if not os.path.exists(backup):
        shutil.copy2(path, backup)
        print("สำรองไฟล์เดิมไว้ที่", backup)

    doc = Document(path)
    print("แพตช์", path)
    toc_index = add_toc_field(doc)
    ensure_toc_styles(doc)
    split_and_number(doc, toc_index)
    doc.save(path)
    print("บันทึกแล้ว:", path)


if __name__ == "__main__":
    target = sys.argv[1] if len(sys.argv) > 1 else os.path.join(
        os.path.dirname(os.path.abspath(__file__)), "user-manual-reserv-2569.docx")
    main(target)
