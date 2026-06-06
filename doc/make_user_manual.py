"""Create user manual DOCX (คู่มือผู้ใช้งาน) for the booking system."""
import os
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

DOC_PATH   = "/mnt/c/projects/reserv/doc/user-manual-v1.docx"
SHOT_DIR   = "/mnt/c/projects/reserv/doc/screenshots"
FONT_NAME  = "TH Sarabun New"
TODAY      = datetime.date.today().strftime("%d %B %Y")

doc = Document()

# ─── Page setup (A4) ──────────────────────────────────────────────────────────
for section in doc.sections:
    section.page_width  = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.0)
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)

# ─── Helpers ──────────────────────────────────────────────────────────────────

def run(para, text, bold=False, size=14, color=None, italic=False):
    r = para.add_run(text)
    r.font.name = FONT_NAME
    r.font.size = Pt(size)
    r.bold   = bold
    r.italic = italic
    if color:
        r.font.color.rgb = RGBColor(*color)
    return r

def para(text="", bold=False, size=14, align=WD_ALIGN_PARAGRAPH.LEFT,
         color=None, space_before=0, space_after=4):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    if text:
        run(p, text, bold=bold, size=size, color=color)
    return p

def heading(text, level=1):
    sizes = {1: 20, 2: 16, 3: 15}
    colors = {1: (0, 100, 0), 2: (6, 78, 59), 3: (34, 68, 34)}
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10 if level == 1 else 6)
    pf.space_after  = Pt(4)
    run(p, text, bold=True, size=sizes.get(level, 14),
        color=colors.get(level, (0, 0, 0)))
    return p

def bullet(text, level=0, size=14):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent  = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.space_before = Pt(1)
    run(p, text, size=size)
    return p

def numbered(text, size=14):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.space_before = Pt(1)
    run(p, text, size=size)
    return p

def add_image(path, width_cm=10, caption=None):
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run()
        r.add_picture(path, width=Cm(width_cm))
        if caption:
            cp = doc.add_paragraph()
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run(cp, caption, size=11, italic=True, color=(100, 100, 100))
    else:
        para(f"[ภาพ: {os.path.basename(path)} — ไม่พบไฟล์]",
             color=(180, 0, 0), italic=True)

def note_box(text, kind="info"):
    colors = {
        "info":    (219, 234, 254),
        "warning": (255, 237, 213),
        "success": (220, 252, 231),
        "danger":  (254, 226, 226),
    }
    text_colors = {
        "info":    (30, 64, 175),
        "warning": (154, 52, 18),
        "success": (21, 128, 61),
        "danger":  (153, 27, 27),
    }
    icons = {"info": "ℹ️", "warning": "⚠️", "success": "✅", "danger": "🚫"}
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.rows[0].cells[0]
    # background
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    bg = colors[kind]
    hex_bg = f"{bg[0]:02X}{bg[1]:02X}{bg[2]:02X}"
    shd.set(qn('w:fill'), hex_bg)
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)
    cp = cell.paragraphs[0]
    cp.paragraph_format.space_before = Pt(4)
    cp.paragraph_format.space_after  = Pt(4)
    run(cp, f"{icons[kind]}  {text}", size=13,
        color=text_colors[kind])
    doc.add_paragraph()  # spacer

def simple_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hrow = t.rows[0]
    for i, h in enumerate(headers):
        c = hrow.cells[i]
        c.paragraphs[0].clear()
        rp = c.paragraphs[0]
        run(rp, h, bold=True, size=13)
        tc = c._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '229799')
        shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)
    for row_data in rows:
        row = t.add_row()
        for i, val in enumerate(row_data):
            c = row.cells[i]
            c.paragraphs[0].clear()
            run(c.paragraphs[0], val, size=13)
    if col_widths:
        n_cols = len(t.columns)
        for ci, w in enumerate(col_widths):
            if ci >= n_cols:
                break
            for r in t.rows:
                r.cells[ci].width = Cm(w)
    doc.add_paragraph()

def page_break():
    doc.add_page_break()

# ─── COVER PAGE ───────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run(p, "\n\n\n", size=14)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run(p, "📚", size=60)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(20)
run(p, "คู่มือการใช้งานระบบ", bold=True, size=28, color=(6, 78, 59))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run(p, "Smart Creative Learning Space", bold=True, size=22, color=(0, 130, 70))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(6)
run(p, "จองพื้นที่บริการ — สำนักวิทยบริการ มหาวิทยาลัยนครพนม", size=16, color=(80, 80, 80))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(8)
run(p, "สำหรับผู้ใช้งานทั่วไป (นักศึกษา / บุคลากร)", size=15, color=(100, 100, 100))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(30)
run(p, f"เวอร์ชัน 1.0  |  {TODAY}", size=13, color=(150, 150, 150))

page_break()

# ─── 1. บทนำ ─────────────────────────────────────────────────────────────────
heading("1. บทนำ", 1)

heading("1.1 ระบบคืออะไร", 2)
para("ระบบจองพื้นที่บริการ Smart Creative Learning Space คือระบบจองห้องออนไลน์ของสำนักวิทยบริการ "
     "มหาวิทยาลัยนครพนม ให้ผู้ใช้บริการสามารถจองห้องสำหรับการเรียนรู้และทำกิจกรรมกลุ่มได้ "
     "ทั้งผ่าน LINE Official Account และผ่านเว็บไซต์")

heading("1.2 ห้องที่ให้บริการ", 2)
simple_table(
    ["ชื่อห้อง", "ความจุ", "สูงสุดต่อครั้ง", "ผู้มีสิทธิ์"],
    [
        ["Mini Theater",          "2–15 คน", "2 ชั่วโมง", "นักศึกษา / บุคลากร"],
        ["Edutainment Room",      "2–10 คน", "2 ชั่วโมง", "นักศึกษา / บุคลากร"],
        ["Canva Design Room",     "1–4 คน",  "2 ชั่วโมง", "นักศึกษา / บุคลากร"],
        ["ChatGPT Room",          "1–4 คน",  "2 ชั่วโมง", "นักศึกษา / บุคลากร"],
        ["ห้องประชุมชั้น 1",      "2–20 คน", "2 ชั่วโมง", "นักศึกษา / บุคลากร"],
    ],
    col_widths=[5, 2.5, 3, 5],
)

heading("1.3 เวลาให้บริการ", 2)
simple_table(
    ["วัน", "เวลาเปิด-ปิด"],
    [
        ["จันทร์ – ศุกร์",   "08:30 – 16:30 น."],
        ["เสาร์ – อาทิตย์",  "09:00 – 17:00 น."],
        ["วันหยุดนักขัตฤกษ์", "ปิดบริการ (แจ้งล่วงหน้า)"],
    ],
    col_widths=[5, 7],
)

note_box("จองล่วงหน้าได้สูงสุด 7 วันทำการ (นับเฉพาะวันเปิดบริการ ไม่รวมวันหยุด)", "info")

page_break()

# ─── 2. ช่องทางการเข้าใช้งาน ────────────────────────────────────────────────
heading("2. ช่องทางการเข้าใช้งาน", 1)

heading("2.1 ผ่าน LINE Official Account (แนะนำ)", 2)
para("วิธีนี้เหมาะสำหรับการใช้งานบนสมาร์ตโฟน ระบบจะยืนยันตัวตนผ่าน LINE โดยอัตโนมัติ")
numbered("เพิ่มเพื่อน LINE OA: Library@NPU")
numbered("กด Rich Menu \"จองห้อง\" ที่ด้านล่างของหน้าแชท")
numbered("เลือกห้องที่ต้องการ และกรอกแบบฟอร์มจอง")

note_box("LINE OA เปิด LIFF ในแอปพลิเคชัน LINE โดยตรง ไม่ต้องล็อกอินเพิ่มเติม", "success")

heading("2.2 ผ่านเว็บไซต์", 2)
para("เปิด browser และเข้าที่ URL:")
para("https://lib.npu.ac.th/reserv/", bold=True, size=14, color=(0, 80, 200))
para("ระบบจะเปิดหน้า LINE Login อัตโนมัติในครั้งแรก กด \"เข้าสู่ระบบด้วย LINE\" แล้วอนุญาตสิทธิ์การเข้าถึง")

note_box("ต้องการ HTTPS เท่านั้น — ไม่รองรับ http:// และไม่รองรับการเปิดในหน้าต่าง Incognito", "warning")

page_break()

# ─── 3. การลงทะเบียนครั้งแรก ────────────────────────────────────────────────
heading("3. การลงทะเบียนครั้งแรก", 1)
para("ผู้ใช้ที่ยังไม่ได้ผูกบัญชี LINE กับระบบมหาวิทยาลัยจะถูก redirect ไปยังหน้าลงทะเบียนโดยอัตโนมัติ "
     "ดำเนินการเพียงครั้งเดียว")

add_image(f"{SHOT_DIR}/user_02_register.png", width_cm=8,
          caption="หน้าลงทะเบียนผูกบัญชี LINE กับระบบ LDAP มหาวิทยาลัย")

heading("ขั้นตอนการลงทะเบียน", 3)
numbered("เลือกประเภทผู้ใช้: \"🎓 นักศึกษา\" หรือ \"👨‍💼 บุคลากร\"")
numbered("กรอก รหัสผู้ใช้ LDAP ได้แก่ รหัสนักศึกษา (เช่น 6401234567) หรือรหัสบุคลากร (เช่น sxxxxxx)")
numbered("กรอก รหัสผ่าน — รหัสเดียวกับระบบ e-Office / อีเมลมหาวิทยาลัย")
numbered("กด \"ลงทะเบียนและเข้าใช้งาน\"")
numbered("ระบบตรวจสอบข้อมูลจากฐานข้อมูล LDAP มหาวิทยาลัย — ใช้เวลาประมาณ 5–10 วินาที")
numbered("เมื่อสำเร็จ ระบบ redirect กลับสู่หน้าที่ต้องการโดยอัตโนมัติ")

note_box("หากลงทะเบียนแล้วเพื่อใช้งานบนเว็บไซต์ไม่จำเป็นต้องลงทะเบียนซ้ำเมื่อใช้ผ่าน LINE OA อีกครั้ง "
         "ระบบจะจำบัญชีที่ผูกไว้แล้ว", "success")

note_box("ข้อผิดพลาดที่พบบ่อย: \"รหัสผ่านไม่ถูกต้อง\" — กรุณาตรวจสอบ Caps Lock และลองพิมพ์ใหม่ "
         "หากลืมรหัสผ่านติดต่อฝ่าย IT มหาวิทยาลัย", "warning")

page_break()

# ─── 4. หน้าหลัก ─────────────────────────────────────────────────────────────
heading("4. หน้าหลัก", 1)

add_image(f"{SHOT_DIR}/user_01_landing.png", width_cm=8,
          caption="หน้าหลัก — แสดงรายการห้องและการจองของฉัน")

heading("ส่วนประกอบของหน้าหลัก", 3)
simple_table(
    ["ส่วน", "คำอธิบาย"],
    [
        ["แถบบน (Header)",          "ชื่อระบบ, ชื่อผู้ใช้ที่ล็อกอิน, ปุ่มไปยังปฏิทิน"],
        ["รายการห้อง",               "แสดงห้องทั้งหมด พร้อมสถานะ ว่าง/ไม่ว่าง คลิกเพื่อเข้าจอง"],
        ["สถานะห้อง (Badge)",        "สีเขียว = ว่าง, สีแดง = ไม่ว่างในขณะนี้"],
        ["การจองของฉัน",             "แสดงการจองที่ยืนยันแล้วของผู้ใช้ในวันนี้และข้างหน้า"],
    ],
    col_widths=[4.5, 11],
)

page_break()

# ─── 5. การจองพื้นที่ ─────────────────────────────────────────────────────────
heading("5. การจองพื้นที่", 1)

heading("5.1 เลือกห้องที่ต้องการ", 2)
para("จากหน้าหลัก คลิกการ์ดห้องที่ต้องการเพื่อเปิดแบบฟอร์มจอง")

add_image(f"{SHOT_DIR}/user_03_booking_form.png", width_cm=8,
          caption="แบบฟอร์มจองห้อง — กรอกข้อมูลการจอง")

heading("5.2 กรอกข้อมูลการจอง", 2)
simple_table(
    ["ฟิลด์", "คำอธิบาย", "ตัวอย่าง"],
    [
        ["ชื่อกลุ่ม / กิจกรรม",  "ชื่อกลุ่มหรือชื่อกิจกรรม",              "ทีมโปรเจกต์ IS301"],
        ["วันที่จอง",             "คลิกเลือกวันจาก date picker",            "5 มิถุนายน 2568"],
        ["เวลาเริ่มต้น",          "เลือก slot เวลาที่ว่าง (สีเทา = ไม่ว่าง)", "10:00"],
        ["ระยะเวลา",              "เลือกระยะเวลาการใช้งาน",                  "1 ชั่วโมง"],
        ["จำนวนผู้เข้าใช้งาน",   "จำนวนสมาชิกกลุ่ม (ต้องอยู่ในช่วงที่ห้องกำหนด)", "5"],
    ],
    col_widths=[4, 6, 5.5],
)

heading("5.3 ยืนยันการจอง", 2)
numbered("ตรวจสอบข้อมูลให้ถูกต้องทั้งหมด")
numbered("กด \"ยืนยันการจอง\"")
numbered("ระบบตรวจสอบช่วงเวลาและความพร้อมของห้อง (ใช้เวลาไม่เกิน 3 วินาที)")
numbered("หากสำเร็จ ระบบแสดงหน้า \"จองสำเร็จแล้ว!\" พร้อมส่ง LINE Notification")

note_box("Slot เวลาสีเทา หมายถึงถูกจองแล้วหรืออยู่นอกเวลาเปิดบริการ ไม่สามารถเลือกได้", "warning")
note_box("จองล่วงหน้าได้สูงสุด 7 วันทำการ นับจากวันที่เปิดบริการ ไม่รวมวันหยุด", "info")

page_break()

# ─── 6. หน้ายืนยันการจอง ────────────────────────────────────────────────────
heading("6. หน้ายืนยันการจอง", 1)

add_image(f"{SHOT_DIR}/user_04_success.png", width_cm=8,
          caption="หน้ายืนยันการจองสำเร็จ")

para("เมื่อจองสำเร็จระบบแสดงหน้ายืนยันพร้อมรายละเอียด ได้แก่ เลขที่การจอง ชื่อห้อง วันที่ เวลา "
     "และชื่อผู้จอง พร้อมกันนั้นระบบส่ง LINE Notification ไปยังบัญชี LINE ของผู้จองโดยอัตโนมัติ")

note_box("หากไม่ได้รับ LINE Notification ภายใน 5 นาที กรุณาตรวจสอบว่าไม่ได้บล็อก LINE OA ของสำนักวิทยบริการ", "info")

page_break()

# ─── 7. การดูการจองของฉัน ────────────────────────────────────────────────────
heading("7. การดูการจองของฉัน", 1)
para("กลับมายังหน้าหลักจะเห็นรายการ \"การจองของฉัน\" ที่ด้านล่าง ซึ่งแสดงการจองที่ยังมีผล (สถานะ: ยืนยัน) "
     "ทั้งวันนี้และวันข้างหน้า")

para("ข้อมูลที่แสดงในแต่ละรายการ:", bold=True)
bullet("ชื่อห้องและไอคอน")
bullet("วันที่และช่วงเวลา")
bullet("สถานะการจอง (ยืนยัน / ถูกยกเลิก)")
bullet("ปุ่มยกเลิกการจอง (กรณีที่ยังไม่ถึงเวลา)")

page_break()

# ─── 8. การยกเลิกการจอง ──────────────────────────────────────────────────────
heading("8. การยกเลิกการจอง", 1)

add_image(f"{SHOT_DIR}/user_05_cancel.png", width_cm=8,
          caption="หน้ายืนยันการยกเลิกการจอง")

numbered("กดปุ่ม \"ยกเลิก\" ที่รายการการจองในหน้าหลัก")
numbered("ระบบแสดง dialog ยืนยัน พร้อมแสดงรายละเอียดการจองที่จะยกเลิก")
numbered("กด \"ยืนยันยกเลิก\" เพื่อดำเนินการ หรือ \"ไม่ยกเลิก\" เพื่อกลับ")
numbered("ระบบบันทึกการยกเลิกและส่ง LINE Notification แจ้งผล")

note_box("การยกเลิกไม่สามารถย้อนกลับได้ หากต้องการใช้งานในช่วงเวลาเดิมอีกครั้ง ต้องจองใหม่", "danger")

page_break()

# ─── 9. การเช็คอิน ───────────────────────────────────────────────────────────
heading("9. การเช็คอิน", 1)

add_image(f"{SHOT_DIR}/user_06_checkin.png", width_cm=8,
          caption="หน้าเช็คอินก่อนเข้าใช้งานห้อง")

para("ก่อนเข้าใช้งานห้องตามเวลาที่จอง ผู้ใช้ต้องเช็คอินผ่านแอปพลิเคชัน")

heading("เงื่อนไขการเช็คอิน", 3)
bullet("เช็คอินได้ตั้งแต่ 15 นาทีก่อนเวลาเริ่มจอง")
bullet("เช็คอินได้ไม่เกิน 15 นาทีหลังเวลาเริ่ม (หลังจากนั้นระบบจะ auto-cancel)")
bullet("กดปุ่ม \"✅ เช็คอิน\" ในหน้าหลัก หรือผ่าน LINE Notification ที่ส่งให้ก่อนถึงเวลา 15 นาที")

note_box("หากไม่เช็คอินภายในเวลาที่กำหนด ระบบจะยกเลิกการจองอัตโนมัติ เพื่อให้ผู้อื่นสามารถจองได้", "warning")

page_break()

# ─── 10. ปฏิทินการจอง (สาธารณะ) ────────────────────────────────────────────
heading("10. ปฏิทินการจอง (สาธารณะ)", 1)
para("ผู้ใช้ทุกคนสามารถดูปฏิทินการจองได้โดยไม่ต้องล็อกอิน ผ่าน URL:")
para("https://lib.npu.ac.th/reserv/calendar/", bold=True, size=14, color=(0, 80, 200))

add_image(f"{SHOT_DIR}/user_09_calendar.png", width_cm=14,
          caption="ปฏิทินแสดงการจองทั้งหมดของทุกห้อง")

bullet("คลิกที่ชื่อเหตุการณ์บนปฏิทินเพื่อดูรายละเอียด")
bullet("กรองตามห้องได้จากเมนู Filter ด้านบน")
bullet("สลับมุมมองระหว่าง รายเดือน / รายสัปดาห์ / รายวัน")

page_break()

# ─── 11. ข้อมูลห้อง (สาธารณะ) ───────────────────────────────────────────────
heading("11. ข้อมูลห้อง", 1)
para("แต่ละห้องมีหน้ารายละเอียดแบบ public ที่ไม่ต้องล็อกอิน สามารถเข้าถึงได้จากปุ่มใน Rich Menu "
     "หรือ URL โดยตรง")

add_image(f"{SHOT_DIR}/user_10_room_detail_mini.png", width_cm=14,
          caption="หน้ารายละเอียดห้อง Mini Theater")

simple_table(
    ["ห้อง", "URL"],
    [
        ["Mini Theater",     "https://lib.npu.ac.th/reserv/room/mini/"],
        ["Edutainment Room", "https://lib.npu.ac.th/reserv/room/edutainment/"],
        ["Canva Design Room","https://lib.npu.ac.th/reserv/room/canva/"],
        ["ChatGPT Room",     "https://lib.npu.ac.th/reserv/room/chat-gpt/"],
        ["ห้องประชุมชั้น 1", "https://lib.npu.ac.th/reserv/room/meeting_f1/"],
    ],
    col_widths=[5, 10.5],
)

page_break()

# ─── 12. บัตรสมาชิกดิจิทัล ──────────────────────────────────────────────────
heading("12. บัตรสมาชิกดิจิทัล (Virtual Card)", 1)

add_image(f"{SHOT_DIR}/user_07_virtual_card.png", width_cm=8,
          caption="บัตรสมาชิกดิจิทัล — แสดงสถานะ Walai@NPU และ QR Code")

para("บัตรสมาชิกดิจิทัลใช้แทนบัตรห้องสมุดจริง สามารถเข้าถึงได้ผ่าน LINE OA หรือ URL:")
para("https://lib.npu.ac.th/reserv/card/", bold=True, size=14, color=(0, 80, 200))

heading("ส่วนประกอบของบัตร", 3)
simple_table(
    ["ส่วน", "คำอธิบาย"],
    [
        ["ชื่อ-นามสกุล",          "ชื่อจริงจากฐานข้อมูลมหาวิทยาลัย"],
        ["ประเภทสมาชิก",          "นักศึกษา หรือ บุคลากรภายในมหาวิทยาลัย"],
        ["สถานะ Walai@NPU",       "แสดงสถานะสมาชิก Walai และจำนวนเล่มที่ยืมได้"],
        ["QR Code",               "ใช้สแกนที่เคาน์เตอร์บริการเพื่อยืนยันตัวตน"],
    ],
    col_widths=[4.5, 11],
)

page_break()

# ─── 13. การควบคุมอุปกรณ์ ────────────────────────────────────────────────────
heading("13. การควบคุมอุปกรณ์ในห้อง (IoT Room Control)", 1)

add_image(f"{SHOT_DIR}/user_08_room_control.png", width_cm=8,
          caption="หน้าควบคุมอุปกรณ์ไฟและแอร์ในห้องที่กำลังใช้งาน")

para("ขณะใช้งานห้องอยู่ ผู้จองสามารถควบคุมอุปกรณ์ภายในห้อง เช่น ไฟ แอร์ โปรเจกเตอร์ ผ่านแอป "
     "ได้โดยตรง โดยไม่ต้องหาสวิตช์")

heading("วิธีเข้าถึง", 3)
numbered("กดปุ่ม \"ควบคุมอุปกรณ์\" บน LINE Notification ที่ส่งเมื่อถึงเวลาจอง")
numbered("หรือเข้าผ่าน URL: https://lib.npu.ac.th/reserv/room-control/")
numbered("ระบบตรวจสอบสิทธิ์ว่าผู้ใช้มีการจองที่ active อยู่ก่อนแสดงหน้าควบคุม")

heading("การเปิด/ปิดอุปกรณ์", 3)
bullet("กดสวิตช์ Toggle ด้านขวาของแต่ละอุปกรณ์เพื่อเปิด/ปิด")
bullet("สถานะสีเขียว = เปิดอยู่, สีเทา = ปิดอยู่")
bullet("หากอุปกรณ์แสดงสถานะ \"ออฟไลน์\" ไม่สามารถควบคุมได้ — แจ้งเจ้าหน้าที่")

note_box("อุปกรณ์จะปิดอัตโนมัติเมื่อหมดเวลาการจอง ระบบส่ง LINE Notification แจ้งก่อน 10 นาที", "info")

page_break()

# ─── 14. คำถามที่พบบ่อย ──────────────────────────────────────────────────────
heading("14. คำถามที่พบบ่อย (FAQ)", 1)

faq = [
    ("จองได้กี่ครั้งต่อวัน?",
     "ไม่มีการจำกัดจำนวนครั้งต่อวัน แต่ไม่สามารถจองซ้อนทับเวลาในห้องเดียวกันได้"),
    ("ลืมเช็คอินทำยังไง?",
     "หากเลยเวลาเช็คอิน 15 นาทีแล้ว ระบบจะยกเลิกการจองอัตโนมัติ ต้องทำการจองใหม่"),
    ("ยกเลิกได้ถึงเมื่อไหร่?",
     "สามารถยกเลิกได้จนถึงก่อนเวลาเริ่มจอง ไม่สามารถยกเลิกหลังถึงเวลาหรือหลังเช็คอินแล้ว"),
    ("ทำไม slot เวลาถึงเป็นสีเทาทั้งหมด?",
     "อาจเกิดจาก (1) วันนั้นมีการปิดห้องชั่วคราว (2) ถูกจองครบแล้ว (3) อยู่นอกเวลาเปิดบริการ "
     "ลองเลือกวันอื่นหรือห้องอื่น"),
    ("เปลี่ยนเวลาการจองได้ไหม?",
     "ระบบยังไม่รองรับการแก้ไขการจอง ต้องยกเลิกแล้วจองใหม่"),
    ("ลงทะเบียนแล้วยังเข้าไม่ได้?",
     "ตรวจสอบว่าบัญชี LINE ที่ใช้ลงทะเบียนตรงกับที่ใช้เปิดแอปในปัจจุบัน หากยังมีปัญหาติดต่อเจ้าหน้าที่"),
    ("บัตรสมาชิกหมดอายุเมื่อไหร่?",
     "บัตรสมาชิกมีอายุตามสถานะนักศึกษา/บุคลากรในฐานข้อมูลมหาวิทยาลัย หากพ้นสภาพจะใช้งานไม่ได้"),
    ("ติดต่อเจ้าหน้าที่ทาง LINE ได้ไหม?",
     "ได้ครับ/ค่ะ สามารถส่งข้อความมาที่ LINE OA Library@NPU ได้ตลอดเวลาทำการ"),
]

for q, a in faq:
    heading(f"Q: {q}", 3)
    para(f"A: {a}")

page_break()

# ─── ภาคผนวก ─────────────────────────────────────────────────────────────────
heading("ภาคผนวก ก — สัญลักษณ์และสถานะ", 1)
simple_table(
    ["สัญลักษณ์ / สี", "ความหมาย"],
    [
        ["🟢 สีเขียว (ว่าง)",       "ห้องพร้อมให้จอง"],
        ["🔴 สีแดง (ไม่ว่าง)",      "ห้องถูกจองแล้วในช่วงเวลานั้น"],
        ["⚫ Slot สีเทา",            "ช่วงเวลานั้นไม่สามารถจองได้ (ถูกจอง / ปิด / นอกเวลา)"],
        ["✅ ยืนยัน",                "การจองสมบูรณ์แล้ว"],
        ["❌ ยกเลิก",                "การจองถูกยกเลิกแล้ว"],
        ["⏰ รอเช็คอิน",             "ถึงเวลาแล้ว รอการยืนยันเข้าใช้"],
    ],
    col_widths=[5.5, 10],
)

heading("ภาคผนวก ข — การแจ้งเตือน LINE", 1)
simple_table(
    ["เหตุการณ์", "เวลาส่ง"],
    [
        ["จองสำเร็จ",                      "ทันทีหลังจอง"],
        ["เตือนก่อนเข้าใช้งาน",            "15 นาทีก่อนเวลาเริ่ม"],
        ["เตือนใกล้หมดเวลา",               "10 นาทีก่อนหมดเวลาการจอง"],
        ["ปิดอุปกรณ์อัตโนมัติ",            "เมื่อหมดเวลาการจอง"],
        ["ยกเลิกโดยเจ้าหน้าที่",           "ทันทีเมื่อถูกยกเลิก"],
        ["ยกเลิกอัตโนมัติ (ไม่เช็คอิน)",   "เมื่อครบ 15 นาทีหลังเวลาเริ่ม"],
    ],
    col_widths=[7, 8.5],
)

heading("ภาคผนวก ค — ข้อมูลติดต่อ", 1)
simple_table(
    ["ช่องทาง", "รายละเอียด"],
    [
        ["LINE OA",    "Library@NPU (เพิ่มเพื่อนในแอป LINE)"],
        ["เว็บไซต์",   "https://lib.npu.ac.th/reserv/"],
        ["โทรศัพท์",   "สำนักวิทยบริการ มหาวิทยาลัยนครพนม"],
        ["เคาน์เตอร์", "ชั้น 1 อาคารสำนักวิทยบริการ ในวันและเวลาทำการ"],
    ],
    col_widths=[3.5, 12],
)

# ─── Save ─────────────────────────────────────────────────────────────────────
doc.save(DOC_PATH)
size_mb = os.path.getsize(DOC_PATH) / (1024 * 1024)
print(f"✓ Saved: {DOC_PATH} ({size_mb:.1f} MB)")
