"""Generate external-access-manual.docx — คู่มือแจ้งเจ้าหน้าที่ เรื่องระบบบุคคลภายนอกเข้าใช้บริการ"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TEAL = RGBColor(0x22, 0x97, 0x99)
DARK = RGBColor(0x1a, 0x1f, 0x2e)
RED = RGBColor(0xdc, 0x35, 0x45)
GRAY = RGBColor(0x71, 0x80, 0x96)
TEAL_BG = RGBColor(0xe6, 0xf4, 0xf4)
FONT = "TH Sarabun New"

doc = Document()

# ── Page setup ────────────────────────────────────────────────────────────────
sec = doc.sections[0]
sec.page_width  = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin    = Cm(2.5)
sec.bottom_margin = Cm(2.5)
sec.left_margin   = Cm(3.0)
sec.right_margin  = Cm(2.0)

# ── Default style ─────────────────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = FONT
style.font.size = Pt(14)
style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT)

def set_heading_style(level, size, color=DARK):
    name = f'Heading {level}'
    s = doc.styles[name]
    s.font.name = FONT
    s.font.size = Pt(size)
    s.font.bold = True
    s.font.color.rgb = color
    s.element.rPr.rFonts.set(qn('w:eastAsia'), FONT)

set_heading_style(1, 20, DARK)
set_heading_style(2, 17, TEAL)
set_heading_style(3, 15, DARK)

def run(para, text, bold=False, size=14, color=None, italic=False, underline=False):
    r = para.add_run(text)
    r.font.name = FONT
    r.font.bold = bold
    r.font.size = Pt(size)
    r.font.italic = italic
    r.font.underline = underline
    r.element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
    if color:
        r.font.color.rgb = color
    return r

def para(text="", bold=False, size=14, color=None, align=WD_ALIGN_PARAGRAPH.LEFT,
         space_before=0, space_after=6, italic=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        run(p, text, bold=bold, size=size, color=color, italic=italic)
    return p

def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    h.paragraph_format.space_after  = Pt(6)
    for r2 in h.runs:
        r2.font.name = FONT
        r2.element.rPr.rFonts.set(qn('w:eastAsia'), FONT)

def bullet(text, indent=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5 + indent * 0.5)
    p.paragraph_format.space_after = Pt(3)
    run(p, text, size=14)

def numbered_list(items, indent=0):
    """Manually-numbered steps (1. 2. 3. ...) that always restart at 1 —
    Word's built-in List Number style shares one counter for the whole
    document, so separate step lists would otherwise keep counting up."""
    for i, text in enumerate(items, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5 + indent * 0.5)
        p.paragraph_format.space_after = Pt(3)
        run(p, f"{i}.  ", bold=True, size=14)
        run(p, text, size=14)

def note_box(text):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = t.cell(0, 0)
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), 'FFF3CD')
    tcPr.append(shd)
    cp = cell.paragraphs[0]
    cp.paragraph_format.space_before = Pt(4)
    cp.paragraph_format.space_after = Pt(4)
    run(cp, "⚠ ", bold=True, size=13, color=RGBColor(0x85, 0x64, 0x04))
    run(cp, text, size=13, color=RGBColor(0x85, 0x64, 0x04))
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def info_box(text):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = t.cell(0, 0)
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), 'E6F4F4')
    tcPr.append(shd)
    cp = cell.paragraphs[0]
    cp.paragraph_format.space_before = Pt(4)
    cp.paragraph_format.space_after = Pt(4)
    run(cp, "ℹ ", bold=True, size=13, color=TEAL)
    run(cp, text, size=13, color=DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def simple_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    hrow = t.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), '1A1F2E')
        tcPr.append(shd)
        cp = cell.paragraphs[0]
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run(cp, h, bold=True, size=12, color=RGBColor(0xFF, 0xFF, 0xFF))
    for ri, row_data in enumerate(rows):
        row = t.rows[ri + 1]
        fill = 'F8F9FA' if ri % 2 == 0 else 'FFFFFF'
        for ci, cell_text in enumerate(row_data):
            cell = row.cells[ci]
            tcPr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), fill)
            tcPr.append(shd)
            cp = cell.paragraphs[0]
            run(cp, str(cell_text), size=12)
    if col_widths:
        for ci, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[ci].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return t

def page_break():
    doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════════════════════════════════
p = para(); p.paragraph_format.space_before = Pt(80)
para("คู่มือแจ้งเจ้าหน้าที่", bold=True, size=32, color=TEAL, align=WD_ALIGN_PARAGRAPH.CENTER)
para("ระบบบุคคลภายนอกเข้าใช้บริการห้องสมุด", bold=True, size=24, color=DARK,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10)
para("(External Access)", size=16, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)
para("ระบบจองพื้นที่บริการ Smart Creative Learning Space", size=15,
     align=WD_ALIGN_PARAGRAPH.CENTER, space_before=20)

p = para(); p.paragraph_format.space_before = Pt(60)
para("สำนักวิทยบริการ มหาวิทยาลัยนครพนม", size=15, align=WD_ALIGN_PARAGRAPH.CENTER)
para("เวอร์ชัน 1.0  |  กรกฎาคม 2569", size=13, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# บทที่ 1 ภาพรวม
# ═══════════════════════════════════════════════════════════════════════════════
heading("บทที่ 1  ภาพรวมระบบ", 1)
para("ระบบบุคคลภายนอกเข้าใช้บริการห้องสมุด เป็นส่วนหนึ่งของระบบจองพื้นที่บริการ Smart Creative "
     "Learning Space ใช้สำหรับออก \"รหัสผ่านเข้าห้องสมุด\" ให้บุคคลภายนอก (ไม่ใช่นักศึกษา/บุคลากรของ "
     "มหาวิทยาลัยที่มีบัญชี LINE ผูกกับระบบ) โดยไม่ต้องผ่านเจ้าหน้าที่ทุกครั้ง")
para("หน้าเว็บทั้งหมดในบทนี้เป็นหน้า public เปิดได้จากลิงก์โดยตรง ไม่ต้อง login และไม่ต้องผ่าน LINE LIFF")

heading("1.1  สองเส้นทางสำหรับบุคคลภายนอก", 2)
para("ระบบแบ่งบุคคลภายนอกออกเป็น 2 กลุ่ม ตามความถี่ในการมาใช้บริการ:")
simple_table(
    ["", "รายวัน (Daily)", "สมาชิกถาวร (Permanent)"],
    [
        ["เหมาะกับใคร", "มาครั้งคราว/ไม่บ่อย", "มาใช้บริการเป็นประจำ, VIP"],
        ["ต้องขอรหัสใหม่ทุกครั้งไหม", "ต้องขอใหม่ทุกวัน", "ขอครั้งเดียว ใช้ได้ตลอด"],
        ["ต้องมีเจ้าหน้าที่อนุมัติไหม", "ไม่ต้อง (ออกรหัสทันที)", "ต้องรออนุมัติจากเจ้าหน้าที่"],
        ["ต้องแนบรูปถ่ายไหม", "ไม่ต้อง", "ต้องแนบตอนลงทะเบียน"],
        ["URL หน้ากรอก", "/reserv/external/", "/reserv/external/permanent/"],
    ],
    col_widths=[4.5, 5, 5.5]
)
info_box("การตรวจสอบเลขบัตรประชาชน (checksum), การอนุมัติ และการออกรหัสทั้งหมด ทำงานอยู่ฝั่ง "
         "api.npu.ac.th ทั้งหมด — เว็บ reserv ทำหน้าที่เป็นเพียงหน้าจอให้กรอกข้อมูลเท่านั้น")

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# บทที่ 2 บุคคลภายนอกรายวัน
# ═══════════════════════════════════════════════════════════════════════════════
heading("บทที่ 2  บุคคลภายนอก \"รายวัน\" — ขอ QR ใช้ได้เฉพาะวันนี้", 1)
para("หน้าเว็บสำหรับผู้ใช้:")
p = para("https://lib.npu.ac.th/reserv/external/", bold=True, size=14, color=TEAL)

heading("2.1  ขั้นตอนที่ผู้ใช้ทำ", 2)
numbered_list([
    "เปิดหน้า /external/ (ไม่ต้อง login)",
    "กรอก ชื่อ-นามสกุล และ เลขบัตรประชาชน 13 หลัก",
    "กดส่งข้อมูล — ระบบตรวจเลขบัตรและออกรหัสให้ทันที (ไม่ต้องรอเจ้าหน้าที่)",
    "ได้รับ QR code หน้าจอเดียวกัน — ใช้เข้าห้องสมุดได้ \"เฉพาะวันนี้\" เท่านั้น",
])

heading("2.2  ข้อความ error ที่ผู้ใช้อาจเจอ (ไว้ตอบคำถามหน้าเคาน์เตอร์)", 2)
simple_table(
    ["สาเหตุ", "ข้อความที่ผู้ใช้เห็น", "สิ่งที่เจ้าหน้าที่ควรทำ"],
    [
        ["เลขบัตรผิด/ไม่ผ่านการตรวจสอบ", "เลขบัตรประชาชนไม่ถูกต้อง กรุณาตรวจสอบแล้วกรอกใหม่",
         "ให้ผู้ใช้ตรวจเลขบัตรอีกครั้ง มักพิมพ์ผิด/สลับหลัก"],
        ["บัญชีถูกระงับสิทธิ์", "บัญชีของคุณถูกระงับสิทธิ์เข้าใช้ กรุณาติดต่อเจ้าหน้าที่",
         "ตรวจสถานะที่ /manage/external/ หากเคยมีประวัติปัญหา"],
        ["โควตารหัสวันนี้เต็ม", "รหัสเข้าใช้สำหรับวันนี้เต็มแล้ว กรุณาติดต่อเจ้าหน้าที่",
         "แจ้งผู้ดูแลระบบ/แจ้งผู้ใช้ให้มาใหม่ในวันถัดไป"],
        ["ระบบ/api ขัดข้องชั่วคราว", "ระบบไม่พร้อมให้บริการขณะนี้ กรุณาลองใหม่หรือติดต่อเจ้าหน้าที่",
         "ลองใหม่อีกครั้งภายหลัง หากยังไม่ได้ให้แจ้งผู้ดูแลระบบ"],
    ],
    col_widths=[4, 6, 5]
)

heading("2.3  การใช้ QR ที่ประตูเข้าห้องสมุด", 2)
para("ทีมประตูสแกนรหัส (access_code 10 หลัก) แล้วระบบฝั่งประตูเรียก endpoint "
     "/v2/external/check/ ของ api.npu.ac.th เพื่อตรวจว่ารหัสนั้นใช้ได้จริงหรือไม่")
note_box("ทีมประตูได้รับตัวอย่าง JSON response ของ endpoint นี้ไปแล้ว (12 ก.ค. 2569) กำลังรอทีมประตู "
         "นำไปทดสอบสแกนจริงหน้างาน — หากพบปัญหาการสแกนไม่ผ่าน ให้แจ้งผู้ดูแลระบบตรวจสอบร่วมกับทีมประตู")

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# บทที่ 3 สมาชิกถาวร (ฝั่งผู้ใช้)
# ═══════════════════════════════════════════════════════════════════════════════
heading("บทที่ 3  บุคคลภายนอก \"สมาชิกถาวร\" — ฝั่งผู้ใช้ (Self-service)", 1)
para("หน้าเว็บสำหรับผู้ใช้:")
p = para("https://lib.npu.ac.th/reserv/external/permanent/", bold=True, size=14, color=TEAL)
para("ใช้สำหรับผู้ที่มาใช้บริการห้องสมุดเป็นประจำ ไม่ต้องขอรหัสใหม่ทุกวัน — เมื่ออนุมัติแล้วจะได้รหัส "
     "ถาวรใบเดียวใช้ได้ตลอด จนกว่าจะถูกระงับ")

heading("3.1  ลงทะเบียนครั้งแรก", 2)
numbered_list([
    "กรอกเลขบัตรประชาชน 13 หลัก (เว้นว่างได้เฉพาะกรณีบุคคลสำคัญ/VIP ที่ไม่มีเลขบัตร เช่น "
    "นายกสภามหาวิทยาลัย — ระบบจะออกรหัสอ้างอิงขึ้นต้นด้วยตัวอักษร V ให้แทน)",
    "ระบบตรวจสอบว่ามีบัตรอยู่แล้วหรือไม่ — ถ้ายังไม่มี จะให้กรอก ชื่อ-นามสกุล และ แนบรูปถ่าย เพิ่มเติม",
    "กดส่งข้อมูล — สถานะจะเป็น \"รออนุมัติ\" (pending) จนกว่าเจ้าหน้าที่จะอนุมัติ (ดูบทที่ 4)",
])

heading("3.2  ตรวจสอบสถานะ / ดูบัตรซ้ำ", 2)
para("ผู้ใช้กลับมากรอกเลขบัตรเดิมที่หน้านี้ได้ทุกเมื่อ เพื่อดูสถานะปัจจุบัน:")
simple_table(
    ["สถานะ", "สิ่งที่ผู้ใช้เห็น"],
    [
        ["ยังไม่เคยลงทะเบียน", "ฟอร์มให้กรอกชื่อ-สกุล + แนบรูป เพื่อลงทะเบียนใหม่"],
        ["pending (รออนุมัติ)", "ข้อความแจ้งว่าอยู่ระหว่างรออนุมัติ"],
        ["active (อนุมัติแล้ว)", "บัตรสมาชิก แสดงรูปถ่าย + QR code ของรหัสถาวร ใช้เข้าห้องสมุดได้ทันที"],
        ["revoked (ถูกระงับ)", "ข้อความแจ้งว่าสิทธิ์ถูกระงับ"],
    ],
    col_widths=[4.5, 10.5]
)
info_box("รูปถ่ายสมาชิกไม่เปิดเป็นลิงก์สาธารณะ — ต้องกรอกเลขบัตรที่ถูกต้องในหน้านี้ หรือเป็นเจ้าหน้าที่ "
         "ที่ login เข้า Staff Portal เท่านั้นจึงจะเห็นรูปได้ (ป้องกันข้อมูลส่วนบุคคลรั่วไหล)")

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# บทที่ 4 สมาชิกถาวร (ฝั่งเจ้าหน้าที่)
# ═══════════════════════════════════════════════════════════════════════════════
heading("บทที่ 4  บุคคลภายนอก \"สมาชิกถาวร\" — ฝั่งเจ้าหน้าที่ (Staff Portal)", 1)
para("เข้าถึงได้จาก Staff Portal หลัง login ที่:")
p = para("https://lib.npu.ac.th/reserv/manage/external/", bold=True, size=14, color=TEAL)

heading("4.1  ดูรายการสมาชิกถาวร", 2)
para("หน้ารายการแสดงสมาชิกถาวรทั้งหมด กรองตามสถานะได้ (รออนุมัติ / อนุมัติแล้ว / ถูกระงับ)")

heading("4.2  ลงทะเบียนแทนผู้ใช้", 2)
para("กรณีผู้ใช้ไม่สะดวกกรอกเอง (เช่น ผู้สูงอายุ, บุคคลสำคัญ) เจ้าหน้าที่ลงทะเบียนแทนได้ที่ "
     "\"+ ลงทะเบียนสมาชิกถาวร\" — กรอกข้อมูลและแนบรูปแทนผู้ใช้ สถานะจะเป็น pending เหมือนผู้ใช้กรอกเอง")

heading("4.3  อนุมัติสมาชิก", 2)
numbered_list([
    "เปิดรายละเอียดสมาชิกที่ต้องการอนุมัติ",
    "ตรวจสอบชื่อ-สกุล และรูปถ่ายให้ตรงกับบัตรประชาชนจริง",
    "กดปุ่ม \"อนุมัติ\" — ระบบจะออกรหัสถาวร (permanent_code) และเปลี่ยนสถานะเป็น active ทันที",
])
note_box("ระบบบันทึกชื่อเจ้าหน้าที่ที่กดอนุมัติจริงไว้เป็นผู้อนุมัติ (approved_by) เพื่อการตรวจสอบย้อนหลัง "
         "— กรุณาตรวจสอบตัวตนให้ถูกต้องก่อนกดอนุมัติทุกครั้ง")

heading("4.4  ระงับสิทธิ์ (Revoke)", 2)
para("ใช้เมื่อพบปัญหา เช่น สมาชิกทำผิดระเบียบ หรือขอยกเลิกสิทธิ์ — กดปุ่ม \"ระงับ\" ที่หน้ารายละเอียด "
     "รหัสถาวรจะใช้งานไม่ได้ทันที โดยข้อมูลยังอยู่ในระบบ สามารถกลับมาตรวจสอบภายหลังได้")

heading("4.5  ลบสมาชิกออกจากระบบ (Delete)", 2)
para("ใช้เฉพาะกรณีลงทะเบียนผิดพลาด หรือไม่ต้องการเก็บข้อมูลไว้แล้วเท่านั้น")
note_box("การลบเป็นการลบถาวร (hard delete) กู้คืนไม่ได้ — หากต้องการเพียงปิดสิทธิ์การใช้งาน ให้ใช้ "
         "\"ระงับ\" (revoke) แทนการลบ")

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# บทที่ 5 คำถามที่พบบ่อย
# ═══════════════════════════════════════════════════════════════════════════════
heading("บทที่ 5  คำถามที่พบบ่อย", 1)

heading("บุคคลภายนอกมาใช้บริการทุกวัน ควรแนะนำให้ขอรหัสรายวันหรือสมาชิกถาวร?", 2)
para("แนะนำให้ลงทะเบียนสมาชิกถาวรที่ /external/permanent/ เพื่อไม่ต้องขอรหัสใหม่ทุกครั้ง "
     "แต่ต้องรอเจ้าหน้าที่อนุมัติก่อนใช้งานได้ครั้งแรก")

heading("QR รายวันใช้ข้ามวันได้ไหม?", 2)
para("ไม่ได้ — รหัสรายวันผูกกับวันที่ออกเท่านั้น หากมาวันถัดไปต้องขอรหัสใหม่ที่ /external/ อีกครั้ง")

heading("สมาชิกถาวรลืมรหัส/QR ต้องทำอย่างไร?", 2)
para("ให้ผู้ใช้กลับไปกรอกเลขบัตรประชาชนเดิมที่ /external/permanent/ อีกครั้ง ระบบจะแสดงบัตร "
     "(รูป + QR) ให้ใหม่ทันทีหากสถานะยังเป็น active — ไม่ต้องลงทะเบียนซ้ำ")

heading("ลงทะเบียนสมาชิกถาวรไปแล้วแต่ยังไม่มีคนอนุมัติ ทำอย่างไร?", 2)
para("แจ้งเจ้าหน้าที่ที่ดูแล Staff Portal ให้เข้าไปตรวจสอบและอนุมัติที่ /manage/external/ "
     "(ดูบทที่ 4.3)")

heading("QR รายวันสแกนที่ประตูไม่ผ่าน ทำอย่างไร?", 2)
para("ตรวจสอบก่อนว่ารหัสยังอยู่ในวันที่ออก (ไม่ข้ามวัน) หากยังไม่ผ่านให้แจ้งผู้ดูแลระบบตรวจสอบร่วมกับ "
     "ทีมประตู เนื่องจากขณะทำเอกสารนี้ การเชื่อมต่อ QR กับระบบประตูยังอยู่ระหว่างทดสอบร่วมกับทีมประตู")

page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# ภาคผนวก
# ═══════════════════════════════════════════════════════════════════════════════
heading("ภาคผนวก — สรุป URL ที่เกี่ยวข้อง", 1)
simple_table(
    ["หน้า", "URL", "ใช้โดยใคร"],
    [
        ["ขอ QR รายวัน", "/reserv/external/", "บุคคลภายนอก (public)"],
        ["ลงทะเบียน/ดูบัตรสมาชิกถาวร", "/reserv/external/permanent/", "บุคคลภายนอก (public)"],
        ["จัดการสมาชิกถาวร", "/reserv/manage/external/", "เจ้าหน้าที่ (login Staff Portal)"],
    ],
    col_widths=[5.5, 6.5, 5]
)

# ── Footer note ───────────────────────────────────────────────────────────────
para()
para("คู่มือนี้จัดทำโดยสำนักวิทยบริการ มหาวิทยาลัยนครพนม",
     size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)
para("Smart Creative Learning Space — ระบบบุคคลภายนอกเข้าใช้บริการ  |  เวอร์ชัน 1.0  |  กรกฎาคม 2569",
     size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)

# ── Save ──────────────────────────────────────────────────────────────────────
out_path = r"C:\projects\reserv\doc\external-access-manual.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
